# reports/views.py
import io
import datetime
import nepali_datetime

from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from django.contrib.auth import get_user_model

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

from duties.models import Duty, DutyChart
from .permissions import IsAdminOrSelf
from users.permissions import user_has_permission_slug, get_allowed_office_ids
import requests

User = get_user_model()


_translation_cache = {}

def translate_to_nepali(text):
    """Google Translate free API for simple names/titles with simple memory cache."""
    if not text or text == "-":
        return text
    
    # Return from cache if available
    if text in _translation_cache:
        return _translation_cache[text]
        
    try:
        url = "https://translate.googleapis.com/translate_a/single"
        params = {"client": "gtx", "sl": "en", "tl": "ne", "dt": "t", "q": text}
        response = requests.get(url, params=params, timeout=5)
        if response.status_code == 200:
            result = response.json()
            if result and result[0] and result[0][0]:
                translated = result[0][0][0]
                _translation_cache[translated] = translated # double save for consistency
                _translation_cache[text] = translated
                return translated
    except Exception as e:
        print(f"Translation error for '{text}': {e}")
        pass
    return text


def translate_to_nepali_batch(text_list):
    """
    Translates a list of strings in batches to minimize API calls.
    Uses \n as a delimiter for the gtx free API.
    """
    if not text_list:
        return {}
    
    # Filter out empty or already cached
    to_translate = [t for t in text_list if t and isinstance(t, str) and str(t) not in _translation_cache]
    if not to_translate:
        return {t: _translation_cache.get(str(t), str(t)) for t in text_list}
    
    # Process in chunks of 50 to stay safe with URL length
    chunk_size = 50
    for i in range(0, len(to_translate), chunk_size):
        chunk = to_translate[i:i + chunk_size]
        combined = "\n".join(chunk)
        
        try:
            url = "https://translate.googleapis.com/translate_a/single"
            params = {
                "client": "gtx",
                "sl": "en",
                "tl": "ne",
                "dt": "t",
                "q": combined
            }
            response = requests.get(url, params=params, timeout=10)
            if response.status_code == 200:
                result = response.json()
                if result and result[0]:
                    # result[0] is typically a list of [[translated_chunk, original_chunk, ...], ...]
                    translated_combined = "".join([part[0] for part in result[0] if part[0]])
                    translated_list = translated_combined.split("\n")
                    
                    # Map back to original strings
                    for orig, trans in zip(chunk, translated_list):
                        _translation_cache[str(orig)] = trans.strip()
        except Exception as e:
            print(f"Batch translation error: {e}")
            
    return {t: _translation_cache.get(str(t), str(t)) for t in text_list}


# ---------------------------
# Helper: Parse user_id[] or comma-separated
# ---------------------------
def _parse_user_ids(request):
    raw = request.GET.getlist("user_id[]") or request.GET.getlist("user_id")
    if not raw:
        s = request.GET.get("user_id")
        if s:
            raw = [s]

    ids = []
    for entry in raw:
        for part in str(entry).split(","):
            try:
                ids.append(int(part.strip()))
            except Exception:
                continue

    # dedupe, preserve order
    seen, out = set(), []
    for i in ids:
        if i not in seen:
            seen.add(i)
            out.append(i)
    return out


# ---------------------------
# Duty options (dropdown)
# ---------------------------
class DutyOptionsView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        qs = DutyChart.objects.all().order_by("effective_date")
        return Response([
            {
                "id": c.id,
                "name": c.name or f"{c.office.name} - {c.effective_date}",
                "effective_date": str(c.effective_date),
                "end_date": str(c.end_date) if c.end_date else str(c.effective_date),
                "office_id": c.office_id,
                "office_name": c.office.name if c.office else "Unknown",
            }
            for c in qs
        ])


# ---------------------------
# Preview JSON (unchanged behavior)
# ---------------------------
class DutyReportPreviewView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")
        if not (date_from and date_to):
            return Response({"groups": []})

        all_users = request.GET.get("all_users") == "1"
        user_ids = [] if all_users else _parse_user_ids(request)
        duty_id = request.GET.get("duty_id")
        schedule_id = request.GET.get("schedule_id")
        office_id = request.GET.get("office_id")

        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")

        qs = Duty.objects.select_related("user", "schedule", "office").filter(
            date__range=[date_from, date_to]
        )

        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        if user_ids:
            qs = qs.filter(user_id__in=user_ids)
        
        if duty_id:
            qs = qs.filter(duty_chart_id=duty_id)

        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)
        
        if office_id and office_id != "all":
            qs = qs.filter(office_id=office_id)

        qs = qs.order_by("user_id", "date", "schedule__start_time")

        groups = {}
        for d in qs:
            if not d.user:
                continue
            uid = d.user.id
            groups.setdefault(uid, {
                "user_id": uid,
                "user_name": getattr(d.user, "full_name", str(d.user)),
                "employee_id": getattr(d.user, "employee_id", "-"),
                "office": d.office.name if d.office else "-",
                "rows": []
            })["rows"].append({
                "id": d.id,
                "date": str(d.date),
                "weekday": d.date.strftime("%A"),
                "schedule": getattr(d.schedule, "name", "-"),
                "start_time": str(d.schedule.start_time) if d.schedule else "-",
                "end_time": str(d.schedule.end_time) if d.schedule else "-",
                "is_completed": d.is_completed,
                "currently_available": d.currently_available,
            })

        return Response({"groups": list(groups.values())})


# ---------------------------
# DOCX EXPORT — अनुसूची-१
# FULL REPLACEMENT (FINAL)
# ---------------------------
class DutyReportFileView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        # -----------------------
        # PARAMS
        # -----------------------
        duty_id = request.GET.get("duty_id")
        all_users = request.GET.get("all_users") == "1"
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")

        # -----------------------
        # FULL CHART MODE
        # -----------------------
        if all_users and duty_id:
            try:
                chart = DutyChart.objects.get(id=duty_id)
                date_from = chart.effective_date
                date_to = chart.end_date or chart.effective_date
            except DutyChart.DoesNotExist:
                return Response({"error": "Invalid duty chart"}, status=400)

        # -----------------------
        # RANGE MODE VALIDATION
        # -----------------------
        if not (date_from and date_to):
            return Response({"error": "Missing dates"}, status=400)

        # -----------------------
        # QUERYSET
        # -----------------------
        qs = Duty.objects.select_related(
            "user", "schedule", "office"
        ).filter(
            date__range=[date_from, date_to]
        )

        if duty_id:
            qs = qs.filter(duty_chart_id=duty_id)

        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")
        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        qs = qs.order_by("date", "schedule__start_time")

        # -----------------------
        # EXCEL GENERATION
        # -----------------------
        fmt = request.GET.get("format", "docx")
        if fmt == "excel":
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Duty Chart"

            # Headers
            headers = [
                "SN", "Designation", "Name", "Phone", 
                "Work Description", "Target", "Timeline", "Remarks"
            ]
            ws.append(headers)

            # Styling Headers
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

            # Data
            sn = 1
            for d in qs:
                row = [
                    sn,
                    (d.user.position.alias or d.user.position.name) if d.user.position else "-",
                    getattr(d.user, "full_name", "-"),
                    getattr(d.user, "phone_number", "-"),
                    "", # Work Description
                    "", # Target
                    str(d.date),
                    ""  # Remarks
                ]
                ws.append(row)
                sn += 1
            
            # Auto-adjust column width (simple)
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter # Get the column name
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column].width = adjusted_width

            bio = io.BytesIO()
            wb.save(bio)
            bio.seek(0)

            return FileResponse(
                bio,
                as_attachment=True,
                filename=f"Duty_Chart_{date_from}_{date_to}.xlsx",
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        # -----------------------
        # DOCX GENERATION
        # -----------------------
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)

        def center(text, bold=False):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = bold
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Header
        center("अनुसूची-१", True)
        center("(परिच्छेद - ३ को दफा ८ र १० सँग सम्बन्धित)")
        center("नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)")
        center("सिफ्ट ड्यूटीमा खटाउनु अघि भर्नु पर्ने बिवरण")

        # Convert dates to Nepali
        try:
            # Handle string or date objects
            d_from = datetime.datetime.strptime(str(date_from), "%Y-%m-%d") if isinstance(date_from, str) else date_from
            d_to = datetime.datetime.strptime(str(date_to), "%Y-%m-%d") if isinstance(date_to, str) else date_to
            
            nepali_from = nepali_datetime.date.from_datetime_date(d_from)
            nepali_to = nepali_datetime.date.from_datetime_date(d_to)
            nepali_period = f"{nepali_from} देखि {nepali_to} सम्म"
        except Exception as e:
            print(f"Failed to convert to Nepali date: {e}")
            nepali_period = f"{date_from} देखि {date_to} सम्म"

        # Unique schedules in this report for classification
        unique_schedules = []
        seen_schedules = set()
        for d in qs:
            if d.schedule and d.schedule.id not in seen_schedules:
                unique_schedules.append(d.schedule)
                seen_schedules.add(d.schedule.id)
        
        def format_time_nepali(t, crosses=False):
            t_str = t.strftime("%H:%M")
            return f"भोलिपल्ट {t_str}" if crosses else t_str

        classification_str = ""
        if unique_schedules:
            parts = []
            for s in unique_schedules:
                crosses = s.end_time < s.start_time
                start_t = s.start_time.strftime("%H:%M")
                end_t = format_time_nepali(s.end_time, crosses)
                parts.append(f"{s.name} ({start_t} — {end_t})")
            classification_str = ", ".join(parts)
        else:
            classification_str = "-"

        # Meta
        meta = doc.add_paragraph()
        meta.add_run("कार्यालयको नाम:- ").bold = True
        meta.add_run(qs.first().office.name if qs.exists() and qs.first().office else "-")

        meta.add_run("\nबिभाग/शाखाको नाम:- ").bold = True
        meta.add_run("\nमिति:- ").bold = True
        meta.add_run(nepali_period)

        meta.add_run("\nड्यूटीको बर्गिकरण:- ").bold = True
        meta.add_run(classification_str)

        doc.add_paragraph("\nकाममा खटाईएको बिवरण:-")

        # Table
        table = doc.add_table(rows=2, cols=8)
        table.style = "Table Grid"

        # Headers Setup with Merging
        # Vertical merges
        table.cell(0, 0).merge(table.cell(1, 0)) # सि.नं.
        table.cell(0, 4).merge(table.cell(1, 4)) # कामको बिवरण
        table.cell(0, 5).merge(table.cell(1, 5)) # लक्ष्य
        table.cell(0, 6).merge(table.cell(1, 6)) # समय सिमा
        table.cell(0, 7).merge(table.cell(1, 7)) # कैफियत

        # Horizontal merge
        table.cell(0, 1).merge(table.cell(0, 3)) # काममा खटाउनु पर्ने कर्मचारीहरुको बिवरण

        # Set Text and Alignment
        def set_cell_text(cell, text, bold=True):
            cell.text = text
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if bold:
                for run in p.runs:
                    run.bold = True

        set_cell_text(table.cell(0, 0), "सि.नं.")
        set_cell_text(table.cell(0, 1), "काममा खटाउनु पर्ने कर्मचारीहरुको विवरण")
        set_cell_text(table.cell(0, 4), "कामको विवरण")
        set_cell_text(table.cell(0, 5), "लक्ष्य")
        set_cell_text(table.cell(0, 6), "समय सिमा")
        
        # Timeline header with timings if single schedule
        timeline_header = "समय सिमा"
        is_single_schedule = unique_schedules and len(unique_schedules) == 1
        if is_single_schedule:
            s = unique_schedules[0]
            crosses = s.end_time < s.start_time
            start_t = s.start_time.strftime("%H:%M")
            end_t = format_time_nepali(s.end_time, crosses)
            timeline_header += f"\n({start_t} - {end_t})"
        set_cell_text(table.cell(0, 6), timeline_header)
        
        set_cell_text(table.cell(0, 7), "कैफियत")

        # Row 1 sub-headers
        set_cell_text(table.cell(1, 1), "पद")
        set_cell_text(table.cell(1, 2), "नाम")
        set_cell_text(table.cell(1, 3), "सम्पर्क नं.")

        # -----------------------
        # BATCH TRANSLATION
        # -----------------------
        unique_texts = set()
        for d in qs:
            if d.user:
                if d.user.full_name: unique_texts.add(d.user.full_name)
                if d.user.position:
                    if d.user.position.alias: unique_texts.add(d.user.position.alias)
                    else: unique_texts.add(d.user.position.name)
        
        translate_to_nepali_batch(list(unique_texts))

        # Helper for Nepali digits
        NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
        def nep(txt):
            return str(txt).translate(NEP_DIGITS)

        sn = 1
        for d in qs:
            cells = table.add_row().cells
            
            # SN
            cells[0].text = nep(f"{sn}.")
            
            # Position
            pos = d.user.position
            pos_nm = (pos.alias or pos.name) if pos else "-"
            cells[1].text = translate_to_nepali(pos_nm)
            
            # Name + Employee ID (in brackets)
            name_str = translate_to_nepali(getattr(d.user, "full_name", "-"))
            emp_id = getattr(d.user, "employee_id", "")
            if emp_id:
                name_str += f" ({nep(emp_id)})"
            cells[2].text = name_str
            
            # Phone
            cells[3].text = nep(getattr(d.user, "phone_number", "-"))
            
            # Work Description & Target
            cells[4].text = ""
            cells[5].text = ""

            # Date / Timeline in Nepali
            d_obj = d.date
            try:
                # Ensure d_obj is a date object
                if isinstance(d_obj, str):
                    d_obj = datetime.datetime.strptime(d_obj, "%Y-%m-%d").date()
                nepali_d = nepali_datetime.date.from_datetime_date(d_obj)
                date_str = nepali_d.strftime("%Y/%m/%d")
            except Exception:
                date_str = str(d_obj)
            
            date_str = nep(date_str)

            if d.schedule:
                if is_single_schedule:
                    cells[6].text = date_str
                else:
                    crosses = d.schedule.end_time < d.schedule.start_time
                    st = d.schedule.start_time.strftime("%H:%M")
                    et = format_time_nepali(d.schedule.end_time, crosses)
                    cells[6].text = f"{date_str}\n({nep(st)} - {nep(et)})"
            else:
                cells[6].text = date_str

            cells[7].text = ""

            # Align cells
            for i in range(8):
                p = cells[i].paragraphs[0]
                # Center most, left align name & work desc
                if i in [2, 4]:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            sn += 1

        # Footer
        doc.add_paragraph(
            "\nकम्पनीको सिफ्ट ड्यूटी निर्देशिका बमोजिम तपाईंहरुलाई माथि उल्लेखित "
            "समयसीमा भित्र कार्य सम्पन्न गर्ने गरी ड्यूटीमा खटाईएको छ | "
            "उक्त कार्य सम्पन्न गरे पश्चात् अनुसूची २ बमोजिम कार्य सम्पन्न गरेको "
            "प्रमाणित गराई पेश गर्नुहुन अनुरोध छ |"
        )

        doc.add_paragraph("\nकाममा खटाउने अधिकार प्राप्त पदाधिकारीको बिवरण:-")
        doc.add_paragraph("नाम:-")
        doc.add_paragraph("पद:-")
        doc.add_paragraph("दस्तखत:-")
        doc.add_paragraph("मिति:-")

        # Response
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        return FileResponse(
            bio,
            as_attachment=True,
            filename=f"Duty_Chart_{date_from}_{date_to}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


class DutyReportNewFileView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        duty_id = request.GET.get("duty_id")
        all_users = request.GET.get("all_users") == "1"
        user_ids = _parse_user_ids(request)
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")
        schedule_id = request.GET.get("schedule_id")
        office_id = request.GET.get("office_id")

        chart = None
        if duty_id:
            try:
                chart = DutyChart.objects.get(id=duty_id)
                # If all_users or if dates weren't provided, use chart dates for filtering
                if all_users or not (date_from and date_to):
                    date_from = chart.effective_date
                    date_to = chart.end_date or chart.effective_date
            except DutyChart.DoesNotExist:
                return Response({"error": "Invalid duty chart"}, status=400)

        if not (date_from and date_to):
            return Response({"error": "Missing dates"}, status=400)

        qs = Duty.objects.select_related("user", "schedule", "office").filter(date__range=[date_from, date_to])
        if not all_users and user_ids:
            qs = qs.filter(user_id__in=user_ids)
        if duty_id:
            qs = qs.filter(duty_chart_id=duty_id)
        
        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)
            
        if office_id and office_id != "all":
            qs = qs.filter(office_id=office_id)
        
        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")
        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        qs = qs.order_by("date", "schedule__start_time")

        doc = Document()
        
        # Moderate margins: Top/Bottom 1", Left/Right 0.75"
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        doc.styles["Normal"].font.size = Pt(11)

        def center(text, bold=False):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = bold
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        center("अनुसूची-२", True)
        center("(परिच्छेद - ३ को दफा ८,९ र १० सँग सम्बन्धित)")
        center("नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)")
        center("सिफ्ट ड्यूटी सम्पन्न भए पश्चात भर्नु पर्ने बिवरण")
        meta = doc.add_paragraph()
        try:
            # If we have a chart, use its official period for the header as requested
            if chart:
                df = chart.effective_date
                dt = chart.end_date or chart.effective_date
            else:
                # Parse dates if they are strings
                if isinstance(date_from, str):
                    df = datetime.datetime.strptime(date_from, "%Y-%m-%d").date()
                else:
                    df = date_from
                    
                if isinstance(date_to, str):
                    dt = datetime.datetime.strptime(date_to, "%Y-%m-%d").date()
                else:
                    dt = date_to
                
            nepali_from = nepali_datetime.date.from_datetime_date(df)
            nepali_to = nepali_datetime.date.from_datetime_date(dt)
            
            # Format: YYYY/MM/DD in Nepali numerals
            nepali_period = f"{str(nepali_from).replace('-', '/')} देखि {str(nepali_to).replace('-', '/')} सम्म"
            NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
            nepali_period = nepali_period.translate(NEP_DIGITS)
        except Exception as e:
            print(f"Date conversion error: {e}")
            nepali_period = f"{date_from} देखि {date_to} सम्म"

        unique_schedules = []
        seen_schedules = set()
        for d in qs:
            if d.schedule and d.schedule.id not in seen_schedules:
                unique_schedules.append(d.schedule)
                seen_schedules.add(d.schedule.id)

        meta = doc.add_paragraph()
        meta.add_run("कार्यालयको नाम:- ").bold = True
        meta.add_run(qs.first().office.name if qs.exists() and qs.first().office else "-")
        meta.add_run("\nबिभाग/शाखाको नाम:- ").bold = True
        meta.add_run("\nमिति:- ").bold = True
        meta.add_run(nepali_period)

        # Duty Classification (बर्गिकरण)
        NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
        def nep(txt): return str(txt).translate(NEP_DIGITS)

        def format_time_nepali(t, crosses=False):
            t_str = t.strftime("%H:%M")
            return f"भोलिपल्ट {t_str}" if crosses else t_str

        parts = []
        for s in unique_schedules:
            crosses = s.end_time < s.start_time
            st = s.start_time.strftime("%H:%M")
            et = format_time_nepali(s.end_time, crosses)
            # Use Nepali digits for times as well
            parts.append(f"{s.name} ({nep(st)} - {nep(et)})")
        
        classification_str = ", ".join(parts) if parts else "-"
        meta.add_run("\nड्यूटीको बर्गिकरण:- ").bold = True
        meta.add_run(classification_str)

        doc.add_paragraph() # Spacer
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run("सम्पन्न भएको कामको बिवरण:-")
        run_desc.bold = True

        table = doc.add_table(rows=2, cols=8)
        table.style = "Table Grid"

        # Adjust Column Widths
        table.columns[0].width = Inches(0.25) # S.N.
        table.columns[1].width = Inches(0.5)  # Position
        table.columns[2].width = Inches(2.45) # Name
        table.columns[3].width = Inches(0.8)  # Contact
        table.columns[4].width = Inches(0.8)  # Work
        table.columns[5].width = Inches(0.6)  # Target
        table.columns[6].width = Inches(1.3)  # Achievement/Time
        table.columns[7].width = Inches(0.6)  # Remarks

        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 6).merge(table.cell(1, 6))
        table.cell(0, 7).merge(table.cell(1, 7))
        table.cell(0, 1).merge(table.cell(0, 3))

        def set_cell_text(cell, text, bold=True):
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if bold:
                for run in p.runs: run.bold = True

        set_cell_text(table.cell(0, 0), "सि.नं.")
        set_cell_text(table.cell(0, 1), "कर्मचारीहरुले सम्पादन गरेको कामको बिवरण")
        set_cell_text(table.cell(0, 4), "कामको विवरण")
        set_cell_text(table.cell(0, 5), "लक्ष्य")
        set_cell_text(table.cell(0, 6), "उपलब्धि")
        set_cell_text(table.cell(0, 7), "काम सम्पादन नभएमा सो को कारण")

        timeline_header = "काम सम्पादन गर्न लागेको समय"
        is_single_schedule = unique_schedules and len(unique_schedules) == 1
        if is_single_schedule:
            s = unique_schedules[0]
            st = s.start_time.strftime("%H:%M")
            et = s.end_time.strftime("%H:%M")
            timeline_header += f"\n({nep(st)} - {nep(et)})"
        set_cell_text(table.cell(1, 6), timeline_header) # achievement subheader essentially

        set_cell_text(table.cell(1, 1), "पद")
        set_cell_text(table.cell(1, 2), "नाम")
        set_cell_text(table.cell(1, 3), "सम्पर्क नं.")

        # -----------------------
        # BATCH TRANSLATION
        # -----------------------
        unique_texts = set()
        for d in qs:
            if d.user:
                if d.user.full_name: unique_texts.add(d.user.full_name)
                if d.user.position:
                    if d.user.position.alias: unique_texts.add(d.user.position.alias)
                    else: unique_texts.add(d.user.position.name)
        
        translate_to_nepali_batch(list(unique_texts))

        NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
        def nep(txt): return str(txt).translate(NEP_DIGITS)

        sn = 1
        for d in qs:
            cells = table.add_row().cells
            cells[0].text = nep(f"{sn}.")
            
            # Position (पद) - Use pre-translated alias/name
            pos = d.user.position
            pos_nm = (pos.alias or pos.name) if pos else "-"
            cells[1].text = translate_to_nepali(pos_nm)
            
            # Name - Use pre-translated full_name
            raw_name = getattr(d.user, "full_name", "-")
            name_str = translate_to_nepali(raw_name)
            emp_id = getattr(d.user, "employee_id", "")
            if emp_id: name_str += f" ({nep(emp_id)})"
            
            cells[2].text = name_str
            cells[3].text = nep(getattr(d.user, "phone_number", "-"))
            cells[4].text = ""
            cells[5].text = ""
            
            d_obj = d.date
            try:
                if isinstance(d_obj, str): d_obj = datetime.datetime.strptime(d_obj, "%Y-%m-%d").date()
                nepali_d = nepali_datetime.date.from_datetime_date(d_obj)
                date_str = nep(str(nepali_d).replace("-", "/"))
            except:
                date_str = nep(str(d_obj).replace("-", "/"))
            cells[6].text = date_str
            cells[7].text = ""

            for i in range(8):
                cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if i in [2, 4] else WD_PARAGRAPH_ALIGNMENT.CENTER
            sn += 1

        
        doc.add_paragraph("")
        doc.add_paragraph("काम सम्पादन भएको प्रमाणित गर्ने पदाधिकारीको विवरण :-")
        
        # Create a 2-column table for signatures
        sig_table = doc.add_table(rows=5, cols=2)
        sig_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Column 1: पेश गर्ने
        sig_table.cell(0, 0).paragraphs[0].add_run("पेश गर्ने:").underline = True
        sig_table.cell(1, 0).text = "नाम :-"
        sig_table.cell(2, 0).text = "पद :-"
        sig_table.cell(3, 0).text = "दस्तखत:-"
        sig_table.cell(4, 0).text = "मिति :-"

        # Column 2: प्रमाणित गर्ने
        sig_table.cell(0, 1).paragraphs[0].add_run("प्रमाणित गर्ने:").underline = True
        sig_table.cell(1, 1).text = "नाम :-"
        sig_table.cell(2, 1).text = "पद :-"
        sig_table.cell(3, 1).text = "दस्तखत:-"
        sig_table.cell(4, 1).text = "मिति :-"

        # Set left indentation for signature fields within columns
        for row in sig_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.left_indent = Inches(0.5)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return FileResponse(bio, as_attachment=True, filename=f"Duty_Report_New_{date_from}_{date_to}.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")