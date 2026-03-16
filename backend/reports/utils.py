# reports/utils.py

import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from datetime import datetime


def generate_report_docx_multi(groups, date_from, date_to):
    """
    groups = [
      {
        user_id,
        user_name,
        employee_id,
        office,
        duties: [Duty objects]
      }
    ]
    """

    doc = Document()

    for idx, g in enumerate(groups):

        if idx > 0:
            doc.add_page_break()

        p = doc.add_paragraph()
        p.add_run("अनुसूची-2\n").bold = True
        p.add_run("(परिच्छेद - ३ को दफा ८, ९ र १० सँग सम्बन्धित)\n")
        p.add_run("नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)\n").bold = True
        p.add_run("सिफ्ट ड्यूटी सम्पन्न भए पश्चात भर्नु पर्ने बिवरण\n\n")

        doc.add_paragraph(f"कार्यालयको नाम:- {g.get('office', '-')}")
        doc.add_paragraph("बिभाग/शाखाको नाम:- Integrated Network Operation Center (iNOC)")
        doc.add_paragraph(f"मिति:- {date_from} देखि {date_to} सम्म")
        doc.add_paragraph("ड्यूटीको बर्गिकरण:-")
        doc.add_paragraph("सम्पन्न भएको कामको बिवरण:-\n")

        table = doc.add_table(rows=1, cols=8)
        hdr = table.rows[0].cells
        headers = [
            "S.N.",
            "पद",
            "नाम",
            "सम्पर्क नं.",
            "कामको बिवरण",
            "लक्ष्य",
            "उपलब्धि",
            "समय",
        ]
        for i, h in enumerate(headers):
            hdr[i].text = h

        for i, duty in enumerate(g["duties"], start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = ""
            row[2].text = g.get("user_name", "")
            row[3].text = ""
            row[4].text = ""
            row[5].text = ""
            row[6].text = ""
            row[7].text = str(getattr(duty.schedule, "end_time", ""))

        doc.add_paragraph("\nकाम सम्पादन भएको प्रमाणित गर्ने पदाधिकारीको बिवरण:-\n")
        doc.add_paragraph("नाम:-")
        doc.add_paragraph("पद:-")
        doc.add_paragraph("दस्तखत:-")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf