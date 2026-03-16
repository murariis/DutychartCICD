# ======================================================================
# FULL views.py (your existing code kept) + UPDATED DutyChartExportFile
# Changes:
#   ✅ PDF export now uses WeasyPrint ONLY (best for Nepali) + FontConfiguration
#   ✅ No xhtml2pdf fallback (it breaks Nepali shaping)
#   ✅ No file:// font-path guessing (use installed fonts in Docker)
#
# NOTE:
#   - You said you will run PDF export from Docker even during Windows dev.
#   - Ensure Docker has fonts + fontconfig (your Dockerfile is already close).
# ======================================================================

from datetime import timedelta
import datetime
import requests
try:
    import nepali_datetime
except ImportError:
    nepali_datetime = None
import os
import platform
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from django.shortcuts import render, get_object_or_404
from django.core.exceptions import ValidationError
from django.db import transaction, IntegrityError

from rest_framework import viewsets, permissions, status, renderers, serializers
from rest_framework.decorators import action
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework.views import APIView
from django.http import HttpResponse, JsonResponse
from django.utils.dateparse import parse_date
from io import BytesIO
import openpyxl

# WeasyPrint import moved to lazy loading (only when PDF export is requested)
# to avoid Windows compatibility issues with GTK libraries
WEASYPRINT_AVAILABLE = False

from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from drf_spectacular.utils import extend_schema, OpenApiExample, OpenApiParameter

from org.models import WorkingOffice as Office
from users.models import User # Added this
from users.permissions import (
    AdminOrReadOnly,
    SuperAdminOrReadOnly,
    IsSuperAdmin,
    IsOfficeAdmin,
    IsOfficeScoped,
    get_allowed_office_ids,
    user_has_permission_slug,
)
from django.core.exceptions import ValidationError, MultipleObjectsReturned # Added this


from .models import DutyChart, Duty, RosterAssignment, Schedule
from .serializers import (
    DutyChartSerializer,
    DutySerializer,
    BulkDocumentUploadSerializer,
    DocumentSerializer,
    ScheduleSerializer,
    ALLOWED_HEADERS,
    HEADER_MAP,
    RosterAssignmentSerializer,
)

import logging
logger = logging.getLogger(__name__)


class ScheduleView(viewsets.ModelViewSet):
    queryset = Schedule.objects.all()
    serializer_class = ScheduleSerializer
    permission_classes = [AdminOrReadOnly]

    @swagger_auto_schema(
        operation_description="List schedules, optionally filtered by office, duty chart, and/or status.",
        manual_parameters=[
            openapi.Parameter("office", openapi.IN_QUERY, description="Filter by Office ID", type=openapi.TYPE_INTEGER),
            openapi.Parameter("duty_chart", openapi.IN_QUERY, description="Filter by Duty Chart ID (schedules linked to chart)", type=openapi.TYPE_INTEGER),
            openapi.Parameter("status", openapi.IN_QUERY, description="Filter by status (e.g., 'template', 'office_schedule')", type=openapi.TYPE_STRING),
        ],
    )
    def list(self, request, *args, **kwargs):
        print(f"[DEBUG] ScheduleView.list() called by user: {request.user}, params: {request.query_params}")
        return super().list(request, *args, **kwargs)

    def get_queryset(self):
        user = self.request.user
        office_id = self.request.query_params.get("office", None)
        duty_chart_id = self.request.query_params.get("duty_chart", None)
        status = self.request.query_params.get("status", None)

        queryset = Schedule.objects.all()

        # 1. STRICT OFFICE FILTERING (High Priority)
        # If an office is selected, we ONLY show schedules for that office.
        # This explicitly EXCLUDES global templates (where office_id is NULL).
        if office_id and str(office_id).strip():
            try:
                if "," in str(office_id):
                    oids = [int(x.strip()) for x in str(office_id).split(",") if x.strip()]
                    queryset = queryset.filter(office_id__in=oids)
                else:
                    oid = int(str(office_id).strip())
                    queryset = queryset.filter(office_id=oid)
            except (ValueError, TypeError):
                return Schedule.objects.none()

        # 2. STATUS & CHART FILTERING
        if status:
            queryset = queryset.filter(status=status)
        if duty_chart_id:
            queryset = queryset.filter(duty_charts__id=duty_chart_id)

        # 3. PERMISSION-BASED VISIBILITY (Access Control)
        # Apply this after functional filters to ensure security.
        if not IsSuperAdmin().has_permission(self.request, self):
            allowed = get_allowed_office_ids(user)
            can_view_any = (
                user_has_permission_slug(user, 'duties.view_any_office_chart') or 
                user_has_permission_slug(user, 'schedules.view_any_office')
            )
            
            if not can_view_any:
                from django.db.models import Q
                
                if office_id:
                    # If they asked for a specific office, ensure they have permission to see it.
                    # Note: We already filtered by office_id above.
                    if int(office_id) not in allowed:
                        return Schedule.objects.none()
                else:
                    # No office specified -> Show global templates + their allowed offices
                    q_res = Q(office_id__isnull=True)
                    if allowed:
                        q_res |= Q(office_id__in=allowed)
                    queryset = queryset.filter(q_res)

        return queryset.distinct().order_by('name')

    def perform_create(self, serializer):
        if IsSuperAdmin().has_permission(self.request, self):
            serializer.save()
            return

        status_val = self.request.data.get("status")
        office_id = self.request.data.get("office")

        # Allow creating global templates (office=null) if status is 'template'
        if status_val == "template" and not office_id:
            if not user_has_permission_slug(self.request.user, 'schedule_templates.create'):
                 raise serializers.ValidationError("You do not have permission to create global schedule templates.")
            serializer.save()
            return

        can_create_any_office = user_has_permission_slug(self.request.user, 'schedules.create_any_office_schedule')

        if not user_has_permission_slug(self.request.user, 'duties.manage_schedule') and \
           not user_has_permission_slug(self.request.user, 'schedules.create') and \
           not can_create_any_office:
            raise serializers.ValidationError("You do not have permission to create schedules.")

        allowed = get_allowed_office_ids(self.request.user)
        if (not office_id or int(office_id) not in allowed) and not can_create_any_office:
            raise serializers.ValidationError("Not allowed to create schedule for this office.")
        serializer.save()


    def perform_update(self, serializer):
        if IsSuperAdmin().has_permission(self.request, self):
            serializer.save()
            return

        status_val = self.request.data.get("status") or getattr(serializer.instance, "status", None)
        office_id = self.request.data.get("office") or getattr(serializer.instance, "office_id", None)

        if status_val == "template" and not office_id:
            if not user_has_permission_slug(self.request.user, 'schedule_templates.edit'):
                 raise serializers.ValidationError("You do not have permission to edit global schedule templates.")
            serializer.save()
            return

        if not user_has_permission_slug(self.request.user, 'duties.manage_schedule'):
            if not user_has_permission_slug(self.request.user, 'schedules.edit'):
                 raise serializers.ValidationError("You do not have permission to update schedules.")

            allowed = get_allowed_office_ids(self.request.user)
            if not office_id or int(office_id) not in allowed:
                raise serializers.ValidationError("Not allowed to update schedule for this office.")
        
        serializer.save()

    def perform_destroy(self, instance):
        if IsSuperAdmin().has_permission(self.request, self):
            instance.delete()
            return

        # Check if it's a global template
        if instance.status == "template" and not instance.office:
             if not user_has_permission_slug(self.request.user, 'schedule_templates.delete'):
                  raise serializers.ValidationError("You do not have permission to delete global schedule templates.")
             instance.delete()
             return

        if not user_has_permission_slug(self.request.user, 'duties.manage_schedule'):
            if not user_has_permission_slug(self.request.user, 'schedules.delete'):
                 raise serializers.ValidationError("You do not have permission to delete schedules.")

            allowed = get_allowed_office_ids(self.request.user)
            if instance.office_id and instance.office_id not in allowed:
                raise serializers.ValidationError("Not allowed to delete schedule for this office.")
        instance.delete()


    @action(detail=False, methods=["post"], url_path="sync-from-roster")
    def sync_from_roster(self, request):
        """
        Pulls all RosterAssignment entries and inserts/updates them into Schedule.
        Dates are ignored; unique schedules are by name+office+start/end times.
        """
        roster_entries = RosterAssignment.objects.all()
        created_count, updated_count = 0, 0

        for ra in roster_entries:
            # Resolve office by name string from roster assignment
            office_obj = None
            if isinstance(ra.office, str) and ra.office:
                office_obj = Office.objects.filter(name__iexact=ra.office.strip()).first()
            elif hasattr(ra, "office") and ra.office and (hasattr(ra.office, "pk") or isinstance(ra.office, int)):
                # If it's a model instance or ID
                if isinstance(ra.office, int):
                    office_obj = Office.objects.filter(id=ra.office).first()
                else:
                    office_obj = ra.office

            if not office_obj:
                # Skip if office cannot be resolved
                continue

            obj, created = Schedule.objects.update_or_create(
                name=ra.shift or "Schedule",
                office=office_obj,
                start_time=ra.start_time,
                end_time=ra.end_time,
                defaults={"status": "active"},
            )
            if created:
                created_count += 1
            else:
                updated_count += 1

        return Response(
            {"message": "Schedule sync complete", "created": created_count, "updated": updated_count},
            status=status.HTTP_200_OK,
        )


class BulkDocumentUploadView(APIView):
    permission_classes = [AdminOrReadOnly]
    parser_classes = [MultiPartParser, FormParser]

    @swagger_auto_schema(
        operation_description="Upload multiple documents in one request.",
        manual_parameters=[
            openapi.Parameter(name="files", in_=openapi.IN_FORM, type=openapi.TYPE_FILE, description="Multiple files to upload", required=True),
            openapi.Parameter(name="meta", in_=openapi.IN_FORM, type=openapi.TYPE_STRING, description="Optional JSON mapping filenames to metadata (e.g. description)", required=False),
        ],
        responses={201: DocumentSerializer(many=True)},
    )
    def post(self, request, *args, **kwargs):
        serializer = BulkDocumentUploadSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        documents = serializer.save()
        return Response(DocumentSerializer(documents, many=True).data, status=status.HTTP_201_CREATED)


class DutyChartViewSet(viewsets.ModelViewSet):
    queryset = DutyChart.objects.all()
    serializer_class = DutyChartSerializer
    permission_classes = [AdminOrReadOnly]

    @swagger_auto_schema(
        operation_description="List duty charts, optionally filtered by office ID.",
        manual_parameters=[
            openapi.Parameter("office", openapi.IN_QUERY, description="Filter by office ID", type=openapi.TYPE_INTEGER)
        ],
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

    def get_queryset(self):
        queryset = DutyChart.objects.all()
        office_id = self.request.query_params.get("office", None)
        user = self.request.user
        if not IsSuperAdmin().has_permission(self.request, self):
            # Check for special 'view_any' permission
            can_view_any = user_has_permission_slug(user, 'duties.view_any_office_chart')
            
            if not can_view_any:
                allowed = get_allowed_office_ids(user)
                if allowed:
                    queryset = queryset.filter(office_id__in=allowed)

        if office_id:
            # Multi-office filter support for dashboard optimization
            oids = []
            if "," in str(office_id):
                try:
                    oids = [int(x.strip()) for x in str(office_id).split(",") if x.strip()]
                except (ValueError, TypeError):
                    return DutyChart.objects.none()
            else:
                try:
                    oids = [int(office_id)]
                except (ValueError, TypeError):
                    return DutyChart.objects.none()

            if not IsSuperAdmin().has_permission(self.request, self):
                can_view_any = user_has_permission_slug(user, 'duties.view_any_office_chart')
                if not can_view_any:
                    allowed = get_allowed_office_ids(user)
                    # Filter oids to only those user is allowed to see
                    oids = [oid for oid in oids if oid in allowed]
                    if not oids:
                        return DutyChart.objects.none()

            queryset = queryset.filter(office_id__in=oids)
        return queryset.order_by("-id")

    def perform_create(self, serializer):
        user = self.request.user
        if IsSuperAdmin().has_permission(self.request, self):
            serializer.save(created_by=user)
            return

        with transaction.atomic():
            office_id = self.request.data.get("office")
            if not office_id:
                raise serializers.ValidationError({"detail": "Office is required."})
            
            # Special case for NETWORK_ADMIN: allow creating for ANY office
            if user.role == 'NETWORK_ADMIN':
                serializer.save(created_by=user)
                return

            allowed = get_allowed_office_ids(user)
            if int(office_id) not in allowed:
                if not user_has_permission_slug(user, 'duties.create_any_office_chart'):
                    raise serializers.ValidationError({"detail": "Not allowed to create duty chart for this office."})
            
            serializer.save(created_by=user)

    def perform_update(self, serializer):
        user = self.request.user

        # 1. SuperAdmin → unrestricted
        if IsSuperAdmin().has_permission(self.request, self):
            serializer.save(edited_by=user)
            return

        # 2. Network Admin → shared access for peers in same office
        if user.role == 'NETWORK_ADMIN':
            chart = serializer.instance
            creator = chart.created_by
            if creator and creator.role == 'NETWORK_ADMIN' and creator.office_id == user.office_id:
                serializer.save(edited_by=user)
                return
        
        # 3. Office-level management: allow if the chart belongs to the user's office
        chart = serializer.instance
        if user.office_id and chart.office_id == user.office_id:
            if user_has_permission_slug(user, 'duties.edit_dutychart'):
                serializer.save(edited_by=user)
                return

        # 4. Other roles → must have edit permission AND must be the creator
        if not user_has_permission_slug(user, 'duties.edit_dutychart'):
            raise serializers.ValidationError({"detail": "You do not have permission to edit duty charts."})

        if chart.created_by_id != user.pk:
            raise serializers.ValidationError({"detail": "You can only edit duty charts that you created or that belong to your office."})

        serializer.save(edited_by=user)

    def perform_destroy(self, instance):
        user = self.request.user

        # 1. SuperAdmin → unrestricted
        if IsSuperAdmin().has_permission(self.request, self):
            instance.delete()
            return

        # 2. Network Admin → shared access for peers in same office
        if user.role == 'NETWORK_ADMIN':
            creator = instance.created_by
            if creator and creator.role == 'NETWORK_ADMIN' and creator.office_id == user.office_id:
                instance.delete()
                return

        # 3. Office-level management: allow if the chart belongs to the user's office
        if user.office_id and instance.office_id == user.office_id:
            if user_has_permission_slug(user, 'duties.delete_chart'):
                instance.delete()
                return

        # 4. Other roles → must have delete permission AND must be the creator
        if not user_has_permission_slug(user, 'duties.delete_chart'):
            raise serializers.ValidationError({"detail": "You do not have permission to delete duty charts."})

        if instance.created_by_id != user.pk:
            raise serializers.ValidationError({"detail": "You can only delete duty charts that you created or that belong to your office."})

        # Safety: cannot delete if employees are still assigned
        if instance.duties.exists():
            raise serializers.ValidationError({
                "detail": "Cannot delete duty chart because employees are assigned to its shifts. Please remove all assignments first."
            })

        instance.delete()


class DutyViewSet(viewsets.ModelViewSet):
    queryset = Duty.objects.all()
    serializer_class = DutySerializer
    permission_classes = [AdminOrReadOnly]

    @swagger_auto_schema(
        operation_description="List duties, optionally filtered by office, user, schedule, and/or date.",
        manual_parameters=[
            openapi.Parameter("office", openapi.IN_QUERY, description="Filter by Office ID", type=openapi.TYPE_INTEGER),
            openapi.Parameter("user", openapi.IN_QUERY, description="Filter by User ID", type=openapi.TYPE_INTEGER),
            openapi.Parameter("schedule", openapi.IN_QUERY, description="Filter by Schedule ID", type=openapi.TYPE_INTEGER),
            openapi.Parameter("date", openapi.IN_QUERY, description="Filter by date (YYYY-MM-DD)", type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE),
            openapi.Parameter("duty_chart", openapi.IN_QUERY, description="Filter by Duty Chart ID", type=openapi.TYPE_INTEGER),
        ],
    )
    def list(self, request, *args, **kwargs):
        return super().list(request, *args, **kwargs)

    def get_queryset(self):
        queryset = Duty.objects.all()
        office_id = self.request.query_params.get("office", None)
        user_id = self.request.query_params.get("user", None)
        schedule_id = self.request.query_params.get("schedule", None)
        date = self.request.query_params.get("date", None)
        duty_chart_id = self.request.query_params.get("duty_chart", None)

        user = self.request.user
        if not IsSuperAdmin().has_permission(self.request, self):
            can_view_any = user_has_permission_slug(user, 'duties.view_any_office_chart')
            if not can_view_any:
                allowed = get_allowed_office_ids(user)
                if allowed:
                    queryset = queryset.filter(office_id__in=allowed)


        if office_id:
            # Multi-office filter support for dashboard optimization
            oids = []
            if "," in str(office_id):
                try:
                    oids = [int(x.strip()) for x in str(office_id).split(",") if x.strip()]
                except (ValueError, TypeError):
                    return Duty.objects.none()
            else:
                try:
                    oids = [int(office_id)]
                except (ValueError, TypeError):
                    return Duty.objects.none()

            if not IsSuperAdmin().has_permission(self.request, self):
                can_view_any = user_has_permission_slug(user, 'duties.view_any_office_chart')
                if not can_view_any:
                    allowed = get_allowed_office_ids(user)
                    # Filter oids to only those user is allowed to see
                    oids = [oid for oid in oids if oid in allowed]
                    if not oids:
                        return Duty.objects.none()
            
            queryset = queryset.filter(office_id__in=oids)

        if user_id:
            queryset = queryset.filter(user_id=user_id)
        if schedule_id:
            queryset = queryset.filter(schedule_id=schedule_id)
        if date:
            queryset = queryset.filter(date=date)
        if duty_chart_id:
            queryset = queryset.filter(duty_chart_id=duty_chart_id)

        return queryset

    def perform_destroy(self, instance):
        user = self.request.user

        # 1. SuperAdmin → unrestricted
        if IsSuperAdmin().has_permission(self.request, self):
            instance.delete()
            return

        # 2. Network Admin management
        if user.role == 'NETWORK_ADMIN':
            duty_chart = instance.duty_chart
            if duty_chart:
                creator = duty_chart.created_by
                if creator and creator.role == 'NETWORK_ADMIN' and creator.office_id == user.office_id:
                    instance.delete()
                    return

        # 3. Office-level management
        duty_chart = instance.duty_chart
        if user.office_id and duty_chart and duty_chart.office_id == user.office_id:
            if user_has_permission_slug(user, 'duties.delete'):
                instance.delete()
                return

        # 4. Must have the base permission to remove employees from duty charts
        if not user_has_permission_slug(user, 'duties.delete'):
            raise serializers.ValidationError("You do not have permission to remove employees from duty charts.")

        # 5. Must be the creator of the duty chart this duty belongs to
        if not duty_chart or duty_chart.created_by_id != user.pk:
            raise serializers.ValidationError(
                "You can only remove employees from duty charts that you created or that belong to your office."
            )

        instance.delete()

    def perform_create(self, serializer):
        user = self.request.user
        if IsSuperAdmin().has_permission(self.request, self):
            serializer.save()
            return

        with transaction.atomic():
            if not user_has_permission_slug(user, 'duties.assign_employee'):
                raise serializers.ValidationError("You do not have permission to assign employees.")

            office_id = self.request.data.get("office")
            chart_id = self.request.data.get("duty_chart")

            # Network Admin special handling
            if user.role == 'NETWORK_ADMIN':
                if chart_id:
                    chart = DutyChart.objects.filter(id=chart_id).first()
                    if chart:
                        creator = chart.created_by
                        if creator and creator.role == 'NETWORK_ADMIN' and creator.office_id == user.office_id:
                            serializer.save()
                            return
            
            # Office Admin / Same Office handling
            if chart_id:
                chart = DutyChart.objects.filter(id=chart_id).first()
                if chart and user.office_id and chart.office_id == user.office_id:
                    if user_has_permission_slug(user, 'duties.assign_employee'):
                        serializer.save()
                        return

            allowed = get_allowed_office_ids(user)
            if office_id is None:
                if chart_id:
                    chart = DutyChart.objects.filter(id=chart_id).first()
                    office_id = getattr(chart, "office_id", None)
            if not office_id or int(office_id) not in allowed:
                raise serializers.ValidationError("Not allowed to create duty for this office.")

            # Check if the assigned user belongs to the same office (unless permission allows otherwise)
            target_user_id = self.request.data.get("user")
            if target_user_id:
                 from users.models import User
                 target_user = User.objects.filter(id=target_user_id).first()
                 
                 if target_user and target_user.office_id:
                      if int(target_user.office_id) != int(office_id):
                          if not user_has_permission_slug(user, 'duties.assign_any_office_employee'):
                              raise serializers.ValidationError(f"You cannot assign employees from other offices ({target_user.office.name}) to this chart.")
            
            serializer.save()


    def perform_update(self, serializer):
        user = self.request.user
        if IsSuperAdmin().has_permission(self.request, self):
            serializer.save()
            return

        with transaction.atomic():
            # Network Admin special handling
            if user.role == 'NETWORK_ADMIN':
                chart = getattr(serializer.instance, "duty_chart", None)
                if chart:
                    creator = chart.created_by
                    if creator and creator.role == 'NETWORK_ADMIN' and creator.office_id == user.office_id:
                        serializer.save()
                        return
            
            # Office Admin / Same Office handling
            chart = getattr(serializer.instance, "duty_chart", None)
            if chart and user.office_id and chart.office_id == user.office_id:
                if user_has_permission_slug(user, 'duties.assign_employee'):
                    serializer.save()
                    return

            allowed = get_allowed_office_ids(user)
            office_id = self.request.data.get("office") or getattr(serializer.instance, "office_id", None)
            if office_id is None:
                chart = getattr(serializer.instance, "duty_chart", None)
                office_id = getattr(chart, "office_id", None)
            if not office_id or int(office_id) not in allowed:
                raise serializers.ValidationError("Not allowed to update duty for this office.")
            serializer.save()

    @swagger_auto_schema(
        method="post",
        operation_description="Bulk create or update duties with shift values.",
        request_body=openapi.Schema(
            type=openapi.TYPE_ARRAY,
            items=openapi.Items(
                type=openapi.TYPE_OBJECT,
                properties={
                    "user": openapi.Schema(type=openapi.TYPE_INTEGER),
                    "office": openapi.Schema(type=openapi.TYPE_INTEGER),
                    "schedule": openapi.Schema(type=openapi.TYPE_INTEGER),
                    "date": openapi.Schema(type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE),
                    "is_completed": openapi.Schema(type=openapi.TYPE_BOOLEAN),
                    "currently_available": openapi.Schema(type=openapi.TYPE_BOOLEAN),
                    "duty_chart": openapi.Schema(type=openapi.TYPE_INTEGER),
                },
            ),
        ),
    )
    @action(detail=False, methods=["post"], url_path="bulk-upsert")
    def bulk_upsert(self, request):
        """
        Bulk create or update duties. Expects list of {user, office, schedule, date, duty_chart, ...}
        """
        try:
            data = request.data
            if not isinstance(data, list):
                raise serializers.ValidationError("Expected a list of duty objects.")

            # Helper to convert to int safely
            def _to_int(val):
                try:
                    if val is None or str(val).strip() == "": return None
                    # Handle float strings like "52.0"
                    if isinstance(val, str) and "." in val:
                        return int(float(val))
                    return int(val)
                except (ValueError, TypeError):
                    return None

            user = request.user
            is_super = IsSuperAdmin().has_permission(request, self)
            
            # 1. Permission checks (Pre-check)
            if not is_super:
                if not user_has_permission_slug(user, 'duties.assign_employee'):
                    raise serializers.ValidationError("You do not have permission to assign employees.")
                
                allowed_offices = get_allowed_office_ids(user)
                is_network_admin = user.role == 'NETWORK_ADMIN'

                chart_cache = {}
                if is_network_admin:
                    c_ids = {_to_int(i.get("duty_chart")) for i in data if i.get("duty_chart")}
                    c_ids.discard(None)
                    if c_ids:
                        charts = DutyChart.objects.filter(id__in=c_ids).select_related('created_by')
                        chart_cache = {c.id: c for c in charts}

                for item in data:
                    cid = _to_int(item.get("duty_chart"))
                    if is_network_admin and cid:
                        chart = chart_cache.get(cid)
                        if chart:
                            creator = chart.created_by
                            if creator and creator.role == 'NETWORK_ADMIN' and creator.office_id == getattr(user, 'office_id', None):
                                continue # Allowed
                            if getattr(user, 'office_id', None) and chart.office_id == user.office_id:
                                continue # Allowed
                        raise serializers.ValidationError(f"Not allowed to assign duty for chart ID {cid}.")

                    oid = _to_int(item.get("office"))
                    if oid is None or oid not in allowed_offices:
                        raise serializers.ValidationError(f"Not allowed to assign duty for office ID {item.get('office')}.")

            # 2. Preparation
            created, updated = 0, 0
            u_ids = {_to_int(i.get("user")) for i in data if i.get("user")}
            u_ids.discard(None)
            s_ids = {_to_int(i.get("schedule")) for i in data if i.get("schedule")}
            s_ids.discard(None)
            c_ids = {_to_int(i.get("duty_chart")) for i in data if i.get("duty_chart")}
            c_ids.discard(None)

            users_map = {u.id: u for u in User.objects.filter(id__in=u_ids)}
            schedules_map = {s.id: s for s in Schedule.objects.filter(id__in=s_ids)}
            charts_map = {c.id: c for c in DutyChart.objects.filter(id__in=c_ids)}

            can_assign_any = is_super or user_has_permission_slug(user, 'duties.assign_any_office_employee')

            # 3. Execution
            with transaction.atomic():
                for item in data:
                    user_id = _to_int(item.get("user"))
                    office_id = _to_int(item.get("office"))
                    chart_id = _to_int(item.get("duty_chart"))
                    schedule_id = _to_int(item.get("schedule"))
                    duty_date_raw = item.get("date")

                    if not user_id or not schedule_id or not duty_date_raw:
                        raise serializers.ValidationError("Missing required fields: user, schedule, or date.")

                    duty_date = parse_date(duty_date_raw) if isinstance(duty_date_raw, str) else duty_date_raw
                    if not duty_date:
                        raise serializers.ValidationError(f"Invalid date format: {duty_date_raw}")

                    if user_id not in users_map:
                        raise serializers.ValidationError(f"User with ID {user_id} does not exist.")
                    if schedule_id not in schedules_map:
                        raise serializers.ValidationError(f"Schedule with ID {schedule_id} does not exist.")
                    if chart_id and chart_id not in charts_map:
                        raise serializers.ValidationError(f"Duty Chart with ID {chart_id} does not exist.")

                    if not can_assign_any:
                        t_user = users_map.get(user_id)
                        t_user_office_id = getattr(t_user, 'office_id', None)
                        if t_user_office_id is None or office_id is None or int(t_user_office_id) != int(office_id):
                            raise serializers.ValidationError(f"Cannot assign employee {t_user.full_name} from a different office (or no office) without the 'Assign Any Office Employee' permission.")

                    obj, was_created = Duty.objects.update_or_create(
                        user_id=user_id,
                        duty_chart_id=chart_id,
                        schedule_id=schedule_id,
                        date=duty_date,
                        defaults={
                            "office_id": office_id,
                            "is_completed": item.get("is_completed", False),
                            "currently_available": item.get("currently_available", True),
                        },
                    )

                    try:
                        obj.full_clean()
                        obj.save()
                    except MultipleObjectsReturned:
                        # This happens if there are duplicate rows for (user, chart, date, schedule)
                        # We try to recover by taking the first one
                        obj = Duty.objects.filter(
                            user_id=user_id, duty_chart_id=chart_id,
                            schedule_id=schedule_id, date=duty_date
                        ).first()
                        obj.office_id = office_id
                        obj.is_completed = item.get("is_completed", False)
                        obj.currently_available = item.get("currently_available", True)
                        obj.save()
                        was_created = False
                    except (ValidationError, IntegrityError) as e:
                        message = getattr(e, 'message_dict', None) or {'detail': str(e)}
                        raise serializers.ValidationError(message)

                    if was_created:
                        created += 1
                    else:
                        updated += 1

            return Response({"created": created, "updated": updated}, status=status.HTTP_200_OK)

        except serializers.ValidationError:
            raise
        except Exception as e:
            logger.exception("Unexpected error in bulk_upsert")
            return Response({"detail": f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @swagger_auto_schema(
        method="post",
        operation_description="Generate a rotation of duties for a user in a date range, cycling shifts.",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            properties={
                "user": openapi.Schema(type=openapi.TYPE_INTEGER),
                "duty_chart": openapi.Schema(type=openapi.TYPE_INTEGER),
                "start_date": openapi.Schema(type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE),
                "end_date": openapi.Schema(type=openapi.TYPE_STRING, format=openapi.FORMAT_DATE),
                "pattern": openapi.Schema(
                    type=openapi.TYPE_ARRAY,
                    items=openapi.Items(type=openapi.TYPE_STRING),
                    description="List of shifts in order to rotate",
                ),
                "overwrite": openapi.Schema(type=openapi.TYPE_BOOLEAN, default=False),
            },
            required=["user", "duty_chart", "start_date", "end_date", "pattern"],
        ),
    )
    @action(detail=False, methods=["post"], url_path="generate-rotation")
    def generate_rotation(self, request):
        user_id = request.data["user"]
        chart_id = request.data["duty_chart"]
        start_date = request.data["start_date"]
        end_date = request.data["end_date"]
        pattern = request.data["pattern"]
        overwrite = request.data.get("overwrite", False)

        if not IsSuperAdmin().has_permission(request, self):
            chart = DutyChart.objects.filter(id=chart_id).first()
            office_id = getattr(chart, "office_id", None)
            allowed = get_allowed_office_ids(request.user)
            if not office_id or int(office_id) not in allowed:
                raise ValidationError("Not allowed to generate rotation for this office.")

        # Check if user is activated
        from users.models import User
        target_user = User.objects.filter(id=user_id).first()
        if not target_user or not target_user.is_activated:
            return Response({"detail": "Cannot generate rotation for a deactivated or non-existent employee."}, status=status.HTTP_400_BAD_REQUEST)

        start = datetime.date.fromisoformat(start_date)
        end = datetime.date.fromisoformat(end_date)
        if end < start:
            return Response({"detail": "end_date must be after or equal to start_date"}, status=status.HTTP_400_BAD_REQUEST)

        days = (end - start).days + 1
        created, updated, skipped = 0, 0, 0

        for i in range(days):
            duty_date = start + timedelta(days=i)
            shift_val = pattern[i % len(pattern)]
            if overwrite:
                obj, was_created = Duty.objects.update_or_create(
                    user_id=user_id,
                    duty_chart_id=chart_id,
                    date=duty_date,
                    defaults={"shift": shift_val},
                )
                if was_created:
                    created += 1
                else:
                    updated += 1
            else:
                obj, was_created = Duty.objects.get_or_create(
                    user_id=user_id,
                    duty_chart_id=chart_id,
                    date=duty_date,
                    defaults={"shift": shift_val},
                )
                if was_created:
                    created += 1
                else:
                    skipped += 1

        return Response({"created": created, "updated": updated, "skipped": skipped}, status=status.HTTP_200_OK)


# ✅ New Roster Bulk Upload View:
class RosterBulkUploadView(APIView):
    permission_classes = [AdminOrReadOnly]
    parser_classes = [MultiPartParser, FormParser]
    querryset = RosterAssignment.objects.all()

    @swagger_auto_schema(
        operation_description=("Bulk upload roster assignments from Excel.\n\n" f"**Required columns:** {', '.join(ALLOWED_HEADERS)}"),
        manual_parameters=[
            openapi.Parameter(
                name="file",
                in_=openapi.IN_FORM,
                type=openapi.TYPE_FILE,
                description=f'Excel file (.xls/.xlsx) with columns: {", ".join(ALLOWED_HEADERS)}',
                required=True,
            )
        ],
        responses={201: "Roster assignments created/updated successfully"},
    )
    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get("file")
        if not file_obj:
            return Response({"detail": "File is required."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file_obj)
        except Exception as e:
            return Response({"detail": f"Invalid Excel file: {e}"}, status=status.HTTP_400_BAD_REQUEST)

        # Normalize headers
        df.columns = [str(c).strip() for c in df.columns]

        # Strict header check
        if list(df.columns) != ALLOWED_HEADERS:
            missing = [c for c in ALLOWED_HEADERS if c not in df.columns]
            extra = [c for c in df.columns if c not in ALLOWED_HEADERS]
            msg_parts = []
            if missing:
                msg_parts.append(f"Missing columns: {', '.join(missing)}")
            if extra:
                msg_parts.append(f"Unexpected columns: {', '.join(extra)}")
            return Response({"detail": " | ".join(msg_parts)}, status=status.HTTP_400_BAD_REQUEST)

        created_count, updated_count, failed_count = 0, 0, 0
        errors = []

        for idx, row in df.iterrows():
            try:
                row_dict = {HEADER_MAP[col]: row[col] for col in ALLOWED_HEADERS}

                # Resolve office FK if needed
                if isinstance(row_dict.get("office"), str):
                    office_obj = Office.objects.filter(name__iexact=row_dict["office"].strip()).first()
                    if not office_obj:
                        failed_count += 1
                        errors.append(f"Row {idx+2}: Working Office '{row_dict['office']}' not found")
                        continue
                    row_dict["office"] = office_obj

                serializer = RosterAssignmentSerializer(data=row_dict)
                serializer.is_valid(raise_exception=True)
                instance = serializer.save()

                # Track created vs updated
                if getattr(instance, "_state", None) and not instance._state.adding:
                    updated_count += 1
                else:
                    created_count += 1

            except Exception as e:
                failed_count += 1
                errors.append(f"Row {idx+2}: {e}")

        detail = f"Created: {created_count}, Updated: {updated_count}, Failed: {failed_count}"
        resp = {"detail": detail}
        if errors:
            resp["errors"] = errors[:10]  # Limit returned errors for safety

        return Response(resp, status=status.HTTP_201_CREATED)


def to_nepali_digits(text):
    if text is None:
        return ""
    text = str(text)
    mapping = {
        '0': '०', '1': '१', '2': '२', '3': '३', '4': '४',
        '5': '५', '6': '६', '7': '७', '8': '८', '9': '९'
    }
    return "".join(mapping.get(c, c) for c in text)


_TRANS_CACHE = {}

def translate_to_nepali(text):
    """
    Translates English text to Nepali using Google Translate free API.
    Uses a simple in-memory cache to minimize redundant network calls.
    """
    if not text or not isinstance(text, str):
        return text
    
    # Check cache first
    if text in _TRANS_CACHE:
        return _TRANS_CACHE[text]
    
    try:
        # Google Translate free API endpoint
        url = "https://translate.googleapis.com/translate_a/single"
        params = {
            "client": "gtx",
            "sl": "en",
            "tl": "ne",
            "dt": "t",
            "q": text
        }
        response = requests.get(url, params=params, timeout=5)
        if response.status_code == 200:
            result = response.json()
            # The result is a nested list: [[["translated_text", "original_text", ...]]]
            if result and result[0] and result[0][0]:
                translated = result[0][0][0]
                _TRANS_CACHE[text] = translated
                return translated
    except Exception as e:
        print(f"Translation error: {e}")
    
    # Return original if translation fails
    return text


def translate_to_nepali_batch(text_list):
    """
    Translates a list of strings in batches to minimize API calls.
    Uses \n as a delimiter for the gtx free API.
    """
    if not text_list:
        return {}
    
    # Filter out empty or already cached
    to_translate = [t for t in text_list if t and isinstance(t, str) and t not in _TRANS_CACHE]
    if not to_translate:
        return {t: _TRANS_CACHE.get(t, t) for t in text_list}
    
    # Process in chunks of 30 to stay safe with URL length (around 2KB)
    chunk_size = 30
    for i in range(0, len(to_translate), chunk_size):
        chunk = to_translate[i:i + chunk_size]
        combined = "\n".join(chunk)
        
        try:
            url = "https://translate.googleapis.com/translate_a/single"
            params = {
                "client": "gtx",
                "sl": "en",
                "tl": "ne",
                "dt": "t",
                "q": combined
            }
            response = requests.get(url, params=params, timeout=10)
            if response.status_code == 200:
                result = response.json()
                if result and result[0]:
                    # result[0] is typically a list of [[translated_chunk, original_chunk, ...], ...]
                    # We join the translated parts and split by \n
                    translated_combined = "".join([part[0] for part in result[0] if part[0]])
                    translated_list = translated_combined.split("\n")
                    
                    # Map back to original strings
                    for orig, trans in zip(chunk, translated_list):
                        _TRANS_CACHE[orig] = trans.strip()
        except Exception as e:
            print(f"Batch translation error: {e}")
            
    return {t: _TRANS_CACHE.get(t, t) for t in text_list}


# ------------------------------------------------------------------------------
# Duty Chart Export: Preview (JSON) and File (Excel/PDF)
# ------------------------------------------------------------------------------

class DutyChartExportPreview(APIView):
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        operation_description=("Preview duty data for a duty chart with optional date range and pagination."),
        manual_parameters=[
            openapi.Parameter("chart_id", openapi.IN_QUERY, description="Duty Chart ID", type=openapi.TYPE_INTEGER, required=True),
            openapi.Parameter("scope", openapi.IN_QUERY, description="Scope: 'full' or 'range'", type=openapi.TYPE_STRING, default="range"),
            openapi.Parameter("start_date", openapi.IN_QUERY, description="Start date (YYYY-MM-DD) when scope=range", type=openapi.TYPE_STRING),
            openapi.Parameter("end_date", openapi.IN_QUERY, description="End date (YYYY-MM-DD) when scope=range", type=openapi.TYPE_STRING),
            openapi.Parameter("page", openapi.IN_QUERY, description="Page number (default 1)", type=openapi.TYPE_INTEGER),
            openapi.Parameter("page_size", openapi.IN_QUERY, description="Items per page (default 50)", type=openapi.TYPE_INTEGER),
            openapi.Parameter("schedule_id", openapi.IN_QUERY, description="Filter by Schedule ID", type=openapi.TYPE_INTEGER),
        ],
    )
    def get(self, request):
        chart_id = request.query_params.get("chart_id")
        scope = (request.query_params.get("scope") or "range").lower()
        start_date_str = request.query_params.get("start_date")
        end_date_str = request.query_params.get("end_date")
        page = int(request.query_params.get("page") or 1)
        page_size = int(request.query_params.get("page_size") or 10)
        schedule_id = request.query_params.get("schedule_id")

        if not chart_id:
            return Response({"detail": "chart_id is required"}, status=status.HTTP_400_BAD_REQUEST)

        chart = get_object_or_404(DutyChart, pk=int(chart_id))
        
        # Optimize: select_related only what is shown in the preview table
        # We don't need directorate, department, or position for the preview grid.
        qs = Duty.objects.filter(duty_chart_id=chart.id).select_related(
            "user", "office", "schedule"
        )

        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)

        if scope == "range":
            if not start_date_str or not end_date_str:
                return Response({"detail": "start_date and end_date are required when scope=range"}, status=status.HTTP_400_BAD_REQUEST)
            start_date = parse_date(start_date_str)
            end_date = parse_date(end_date_str)
            if not start_date or not end_date:
                return Response({"detail": "Invalid start_date or end_date"}, status=status.HTTP_400_BAD_REQUEST)
            qs = qs.filter(date__gte=start_date, date__lte=end_date)

        # Remove qs.count() to save a database trip as it's not used in the UI
        # Pagination
        offset = (page - 1) * page_size
        items = list(qs.order_by("date")[offset : offset + page_size])

        def duty_to_row(d: Duty):
            user = d.user
            office = d.office
            schedule = d.schedule
            return {
                "date": d.date.isoformat(),
                "employee_id": getattr(user, "employee_id", None),
                "full_name": getattr(user, "full_name", None) or getattr(user, "username", None),
                "position": getattr(user.position, "name", "-") if user and user.position else "-",
                "phone_number": getattr(user, "phone_number", None),
                "office": getattr(office, "name", None) or (getattr(getattr(user, "office", None), "name", None) if user else None),
                "schedule": getattr(schedule, "name", None),
                "start_time": getattr(schedule, "start_time", None).strftime("%H:%M") if getattr(schedule, "start_time", None) else None,
                "end_time": getattr(schedule, "end_time", None).strftime("%H:%M") if getattr(schedule, "end_time", None) else None,
            }

        rows = [duty_to_row(d) for d in items]

        payload = {
            "chart": {
                "id": chart.id,
                "name": chart.name,
                "office": getattr(chart.office, "name", None),
                "effective_date": chart.effective_date.isoformat(),
                "end_date": chart.end_date.isoformat() if chart.end_date else None,
            },
            "columns": [
                {"key": "date", "label": "Date"},
                {"key": "employee_id", "label": "Employee ID"},
                {"key": "full_name", "label": "Employee Name"},
                {"key": "position", "label": "Position"},
                {"key": "phone_number", "label": "Phone"},
                {"key": "office", "label": "Office"},
                {"key": "schedule", "label": "Schedule"},
                {"key": "start_time", "label": "Start Time"},
                {"key": "end_time", "label": "End Time"},
            ],
            "rows": rows,
        }

        return Response(payload, status=status.HTTP_200_OK)


# ------------------------------------------------------------------------------
# Content Negotiation helper (keep your existing behavior)
# ------------------------------------------------------------------------------
from rest_framework.negotiation import DefaultContentNegotiation

class IgnoreFormatContentNegotiation(DefaultContentNegotiation):
    def get_format_suffix(self, request):
        return None  # Ignore the 'format' query parameter suffix


# ------------------------------------------------------------------------------
# UPDATED: DutyChartExportFile (Excel/PDF/DOCX)
#   ✅ PDF = WeasyPrint only + Nepali font family names
#   ✅ Works reliably in Linux/Docker (your intended runtime)
# ------------------------------------------------------------------------------
class DutyChartExportFile(APIView):
    permission_classes = [permissions.IsAuthenticated]
    content_negotiation_class = IgnoreFormatContentNegotiation

    def get(self, request):
        try:
            return self._get(request)
        except Exception as e:
            import traceback
            error_msg = f"EXPORT ERROR: {str(e)}\n{traceback.format_exc()}"
            with open("/home/subin/duty-chart/backend/export_error.log", "a") as f:
                f.write(error_msg + "\n" + "="*40 + "\n")
            return Response({"detail": f"Export failed: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def _get(self, request):
        chart_id = request.query_params.get("chart_id")
        out_format = (request.query_params.get("export_format") or request.query_params.get("format") or "").lower()
        scope = (request.query_params.get("scope") or "range").lower()
        start_date_str = request.query_params.get("start_date")
        end_date_str = request.query_params.get("end_date")
        schedule_id = request.query_params.get("schedule_id")

        if not chart_id:
            return Response({"detail": "chart_id is required"}, status=status.HTTP_400_BAD_REQUEST)

        chart = get_object_or_404(DutyChart, pk=int(chart_id))
        qs = Duty.objects.filter(duty_chart_id=chart.id).select_related(
            "user", "office", "schedule", "user__directorate", "user__department", "user__position"
        )

        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)

        if scope == "range":
            if not start_date_str or not end_date_str:
                return Response({"detail": "start_date and end_date are required when scope=range"}, status=status.HTTP_400_BAD_REQUEST)
            start_date = parse_date(start_date_str)
            end_date = parse_date(end_date_str)
            if not start_date or not end_date:
                return Response({"detail": "Invalid start_date or end_date"}, status=status.HTTP_400_BAD_REQUEST)
            qs = qs.filter(date__gte=start_date, date__lte=end_date)

        # Prepare data for headers
        if scope == "range":
            final_start = start_date
            final_end = end_date
        else:
            final_start = chart.effective_date
            final_end = chart.end_date or chart.effective_date

        nepali_period = f"{to_nepali_digits(str(final_start).replace('-', '/'))} देखि {to_nepali_digits(str(final_end).replace('-', '/'))}"
        if nepali_datetime:
            try:
                nep_start = nepali_datetime.date.from_datetime_date(final_start)
                nep_end = nepali_datetime.date.from_datetime_date(final_end)
                nepali_period = f"{to_nepali_digits(str(nep_start).replace('-', '/'))} देखि {to_nepali_digits(str(nep_end).replace('-', '/'))} सम्म"
            except:
                pass

        unique_schedules = []
        seen_sch = set()
        for d in qs:
            if d.schedule and d.schedule.id not in seen_sch:
                unique_schedules.append(d.schedule)
                seen_sch.add(d.schedule.id)
        
        if unique_schedules:
            class_parts = []
            for s in unique_schedules:
                st = s.start_time.strftime("%H:%M")
                et = s.end_time.strftime("%H:%M")
                class_parts.append(f"{s.name} ({st} - {et})")
            classification = ", ".join(class_parts)
        else:
            classification = getattr(chart, "name", "-") or "-"

        rows = []
        for d in qs.order_by("date"):
            user = d.user
            office = d.office
            schedule = d.schedule
            pos = getattr(user, "position", None) if user else None

            # Timing for individual row
            timing = ""
            if schedule:
                st = schedule.start_time.strftime("%H:%M")
                et = schedule.end_time.strftime("%H:%M")
                timing = f"\n({st} - {et})"

            rows.append(
                [
                    f"{d.date.isoformat()}{timing}",
                    getattr(user, "employee_id", "") or "",
                    (getattr(user, "full_name", "") or getattr(user, "username", "")) if user else "",
                    getattr(user, "phone_number", "") if user else "",
                    getattr(getattr(user, "directorate", None), "name", "") if user else "",
                    getattr(getattr(user, "department", None), "name", "") if user else "",
                    getattr(office, "name", "") or (getattr(getattr(user, "office", None), "name", "") if user else ""),
                    getattr(schedule, "name", "") or "",
                    getattr(schedule, "start_time", None).strftime("%H:%M") if getattr(schedule, "start_time", None) else "",
                    getattr(schedule, "end_time", None).strftime("%H:%M") if getattr(schedule, "end_time", None) else "",
                    getattr(pos, "alias", None) or getattr(pos, "name", "-") if pos else "-",
                ]
            )

        # Pre-translate unique names to avoid sequential API call bottleneck
        unique_names = list(set(r[2] for r in rows if r[2]))
        translate_to_nepali_batch(unique_names)

        # ---------------- Excel ----------------
        if out_format == "excel":
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Duty Export"
            headers = ["Date", "Employee ID", "Employee Name", "Phone", "Directorate", "Department", "Office", "Schedule", "Start Time", "End Time", "Position"]
            ws.append(headers)
            
            # Bold headers
            from openpyxl.styles import Font
            for cell in ws[1]:
                cell.font = Font(bold=True)

            for r in rows:
                ws.append(r)
            bio = BytesIO()
            wb.save(bio)
            bio.seek(0)
            filename = f"duty_chart_{chart.id}_{datetime.date.today().isoformat()}.xlsx"
            resp = HttpResponse(bio.getvalue(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            resp["Content-Disposition"] = f'attachment; filename="{filename}"'
            return resp

        # ---------------- PDF (WeasyPrint ONLY) ----------------
        if out_format == "pdf":
            # Lazy import so Windows host doesn't need GTK (you will run this inside Docker anyway)
            try:
                from weasyprint import HTML, CSS
                from weasyprint.text.fonts import FontConfiguration
            except Exception as e:
                return Response(
                    {
                        "detail": (
                            "WeasyPrint is not available in this runtime. "
                            "Run PDF export inside Docker (Linux image) where WeasyPrint deps are installed. "
                            f"Error: {str(e)}"
                        )
                    },
                    status=status.HTTP_501_NOT_IMPLEMENTED,
                )

            headers_np = ["सि.नं.", "पद", "नाम", "सम्पर्क नं.", "कामको बिवरण", "लक्ष्य", "समय सिमा", "कैफियत"]

            table_rows_html = ""
            for idx, r in enumerate(rows, start=1):
                # Columns: SN, Position, Name(ID), Phone, WorkDesc, Target, Timeline, Remarks
                sn_np = to_nepali_digits(idx)
                pos = r[10] # পদ stays in English
                translated_name = translate_to_nepali(r[2])
                name_np = f"{translated_name} ({to_nepali_digits(r[1])})" # Name (ID in Nepali)
                phone_np = to_nepali_digits(r[3])
                
                # Column 6: Date handling (Convert to Nepali digits)
                just_date = r[0].split('\n')[0]
                if nepali_datetime:
                   try:
                       d_obj = datetime.date.fromisoformat(just_date)
                       n_date = nepali_datetime.date.from_datetime_date(d_obj)
                       just_date = str(n_date)
                   except:
                       pass
                timeline_np = to_nepali_digits(just_date.replace('-', '/'))

                row_data = [sn_np, pos, name_np, phone_np, "", "", timeline_np, ""]
                table_rows_html += "<tr>" + "".join([f"<td>{cell}</td>" for cell in row_data]) + "</tr>"

            html_str = f"""
            <html>
            <head>
                <meta charset="utf-8" />
            </head>
            <body>
                <div class="center header">
                    <div class="bold">अनुसूची-१</div>
                    <div>(परिच्छेद - ३ को दफा ८ र १० सँग सम्बन्धित)</div>
                    <div class="bold title">नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)</div>
                    <div class="bold">सिफ्ट ड्युटीमा खटाउनु अघि भर्नु पर्ने बिवरण</div>
                </div>

                <div class="meta">
                    <div><strong>कार्यालयको नाम:-</strong> {getattr(chart.office, 'name', '-')}</div>
                    <div><strong>बिभाग/शाखाको नाम:-</strong> </div>
                    <div><strong>मिति:-</strong> {nepali_period}</div>
                    <div><strong>ड्युटीको बर्गिकरण:-</strong> {classification}</div>
                </div>

                <table>                    <thead>
                        <tr>{"".join([f"<th>{h}</th>" for h in headers_np])}</tr>
                    </thead>
                    <tbody>
                        {table_rows_html}
                    </tbody>
                </table>

                <div class="note">
                    कम्पनीको सिफ्ट ड्युटी निर्देशिका बमोजिम तपाईंहरुलाई माथि उल्लेखित समय सीमा भित्र कार्य सम्पन्न गर्ने गरी ड्युटीमा खटाईएको छ |
                    उक्त कार्य सम्पन्न गरे पश्चात् अनुसूची-२ बमोजिम कार्य सम्पन्न गरेको प्रमाणित गराई पेश गर्नुहुन अनुरोध छ |
                </div>

                <div class="sign-title">काममा खटाउने अधिकार प्राप्त पदाधिकारीको विवरण :-</div>
                <div class="sign-container">
                    <div class="sign-block">
                        <div class="u">सिफारिस गर्ने:</div>
                        <div>नाम :-</div>
                        <div>पद :-</div>
                        <div>दस्तखत:-</div>
                        <div>मिति :-</div>
                    </div>
                    <div class="sign-block">
                        <div class="u">स्वीकृत गर्ने:</div>
                        <div>नाम :-</div>
                        <div>पद :-</div>
                        <div>दस्तखत:-</div>
                        <div>मिति :-</div>
                    </div>
                </div>
            </body>
            </html>
            """

            css_str = """
            @page { 
                size: A4; 
                margin-top: 1in; 
                margin-bottom: 1in; 
                margin-left: 0.75in; 
                margin-right: 0.75in; 
            }
            body {
                font-family: "Noto Sans Devanagari", "Nirmala UI", "Mangal", "DejaVu Sans", sans-serif;
                font-size: 10pt;
                line-height: 1.4;
            }
            .center { text-align: center; }
            .bold { font-weight: bold; }
            .header { margin-bottom: 25px; }
            .title { font-size: 14pt; }
            .meta { margin-bottom: 20px; }
            table { width: 100%; border-collapse: collapse; margin-top: 10px; }
            th, td { border: 1px solid #000; padding: 5px; text-align: left; vertical-align: top; }
            th:nth-child(1), td:nth-child(1) { width: 5%; text-align: center; } /* S.N. */
            th:nth-child(2), td:nth-child(2) { width: 14%; } /* Position */
            th:nth-child(3), td:nth-child(3) { width: 22%; } /* Name */
            th:nth-child(4), td:nth-child(4) { width: 10%; } /* Phone */
            th:nth-child(5), td:nth-child(5) { width: 12%; } /* Work */
            th:nth-child(6), td:nth-child(6) { width: 10%; } /* Target */
            th:nth-child(7), td:nth-child(7) { width: 19%; white-space: nowrap; } /* Timeline */
            th:nth-child(8), td:nth-child(8) { width: 8%; }  /* Remarks */
            th { background: #f2f2f2; font-weight: bold; }
            .note { margin-top: 25px; }
            .sign-title { margin-top: 25px; font-weight: bold; }
            .sign-container { margin-top: 10px; display: flex; justify-content: space-between; }
            .sign-block { width: 45%; line-height: 1.8; }
            .u { text-decoration: underline; font-weight: bold; }
            """

            try:
                font_config = FontConfiguration()
                pdf_bytes = HTML(string=html_str, base_url=".").write_pdf(
                    stylesheets=[CSS(string=css_str, font_config=font_config)],
                    font_config=font_config,
                )
            except Exception as e:
                return Response({"detail": f"PDF generation failed: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            filename = f"duty_chart_{chart.id}_{datetime.date.today().isoformat()}.pdf"
            resp = HttpResponse(pdf_bytes, content_type="application/pdf")
            resp["Content-Disposition"] = f'attachment; filename="{filename}"'
            return resp

        # ---------------- DOCX ----------------
        if out_format == "docx":
            doc = Document()
            
            # Moderate margins: Top/Bottom 1", Left/Right 0.75"
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            style = doc.styles["Normal"]
            style.font.size = Pt(11)

            p = doc.add_paragraph("अनुसूची-१\n(परिच्छेद - ३ को दफा ८ र १० सँग सम्बन्धित)")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].bold = True

            p = doc.add_paragraph("नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            p = doc.add_paragraph("सिफ्ट ड्युटीमा खटाउनु अघि भर्नु पर्ने बिवरण")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("")

            meta = doc.add_paragraph()
            meta.add_run("कार्यालयको नाम:- ").bold = True
            meta.add_run(getattr(chart.office, "name", "-"))
            meta.add_run("\n")

            meta.add_run("बिभाग/शाखाको नाम:- ").bold = True
            meta.add_run("\n")

            meta.add_run("मिति:- ").bold = True
            meta.add_run(nepali_period)
            meta.add_run("\n")

            meta.add_run("ड्यूटीको बर्गिकरण:- ").bold = True
            meta.add_run(classification)
            meta.add_run("\n")
            meta.add_run("\n")

            meta.add_run("काममा खटाईएको बिवरण:- ").bold = True
            # Table Header as per image
            table = doc.add_table(rows=2, cols=8)
            table.style = "Table Grid"

            # Merge सि.नं.
            c0 = table.cell(0, 0)
            c0.merge(table.cell(1, 0))
            c0.text = "सि.नं."

            # Merge काममा खटाउनु पर्ने कर्मचारीहरुको बिवरण
            c1_3 = table.cell(0, 1)
            c1_3.merge(table.cell(0, 3))
            c1_3.text = "काममा खटाउनु पर्ने कर्मचारीहरुको बिवरण"

            table.cell(1, 1).text = "पद"
            table.cell(1, 2).text = "नाम"
            table.cell(1, 3).text = "सम्पर्क नं."

            # Merge कामको बिवरण
            c4 = table.cell(0, 4)
            c4.merge(table.cell(1, 4))
            c4.text = "कामको बिवरण"

            # Merge लक्ष्य
            c5 = table.cell(0, 5)
            c5.merge(table.cell(1, 5))
            c5.text = "लक्ष्य"

            # Merge समय सिमा
            c6 = table.cell(0, 6)
            c6.merge(table.cell(1, 6))
            time_header = "समय सिमा"
            if unique_schedules and len(unique_schedules) == 1:
                s = unique_schedules[0]
                st = s.start_time.strftime("%H:%M")
                et = s.end_time.strftime("%H:%M")
                time_range = f"\n({st} - "
                if s.end_time < s.start_time:
                    time_range += f"भोलिपल्ट {et})"
                else:
                    time_range += f"{et})"
                time_header += time_range
            c6.text = to_nepali_digits(time_header)

            # Merge कैफियत
            c7 = table.cell(0, 7)
            c7.merge(table.cell(1, 7))
            c7.text = "कैफियत"

            # Adjust Column Widths (A4 width is approx 7 in)
            table.columns[0].width = Inches(0.4)  # S.N.
            table.columns[1].width = Inches(1.1)  # Position (Left in English)
            table.columns[2].width = Inches(1.6)  # Name
            table.columns[3].width = Inches(0.8)  # Phone
            table.columns[4].width = Inches(0.8)  # Work
            table.columns[5].width = Inches(0.6)  # Target
            table.columns[6].width = Inches(1.3)  # Timeline (Increased to avoid wrap)
            table.columns[7].width = Inches(0.6)  # Remarks

            # Styling headers
            for r_idx in range(2):
                for c_idx in range(8):
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in p.runs:
                            run.bold = True

            # Data Rows
            for idx, r in enumerate(rows, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = to_nepali_digits(idx)
                row_cells[1].text = r[10]  # Position (Left in English as requested)
                translated_name = translate_to_nepali(r[2])
                row_cells[2].text = f"{translated_name} ({to_nepali_digits(r[1])})"  # Name (ID in Nepali)
                row_cells[3].text = to_nepali_digits(r[3])   # Phone in Nepali
                row_cells[4].text = ""     # Work desc
                row_cells[5].text = ""     # Target

                # Column 6: Date
                just_date = r[0].split('\n')[0]
                if nepali_datetime:
                   try:
                       d_obj = datetime.date.fromisoformat(just_date)
                       n_date = nepali_datetime.date.from_datetime_date(d_obj)
                       just_date = str(n_date)
                   except:
                       pass
                row_cells[6].text = to_nepali_digits(just_date.replace('-', '/'))
                row_cells[7].text = ""

                # Align data rows
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("")
            footer_msg = doc.add_paragraph(
                "कम्पनीको सिफ्ट ड्युटी निर्देशिका बमोजिम तपाईंहरुलाई माथि उल्लेखित समय सीमा भित्र कार्य सम्पन्न गर्ने गरी ड्युटीमा खटाईएको छ | "
                "उक्त कार्य सम्पन्न गरे पश्चात् अनुसूची-२ बमोजिम कार्य सम्पन्न गरेको प्रमाणित गराई पेश गर्नुहुन अनुरोध छ |"
            )
            footer_msg.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            doc.add_paragraph("")
            doc.add_paragraph("काममा खटाउने अधिकार प्राप्त पदाधिकारीको विवरण :-")
            
            # Create a 2-column table for signatures
            sig_table = doc.add_table(rows=5, cols=2)
            sig_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Column 1: सिफारिस गर्ने
            sig_table.cell(0, 0).paragraphs[0].add_run("सिफारिस गर्ने:").underline = True
            sig_table.cell(1, 0).text = "नाम :-"
            sig_table.cell(2, 0).text = "पद :-"
            sig_table.cell(3, 0).text = "दस्तखत:-"
            sig_table.cell(4, 0).text = "मिति :-"

            # Column 2: स्वीकृत गर्ने
            sig_table.cell(0, 1).paragraphs[0].add_run("स्वीकृत गर्ने:").underline = True
            sig_table.cell(1, 1).text = "नाम :-"
            sig_table.cell(2, 1).text = "पद :-"
            sig_table.cell(3, 1).text = "दस्तखत:-"
            sig_table.cell(4, 1).text = "मिति :-"

            # Set alignment for all signature cells to match the image (mostly left within their columns)
            # but we can indent the columns as needed.
            for row in sig_table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        # Add some left padding/indent to make it look like the image
                        p.paragraph_format.left_indent = Inches(0.5)

            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)

            filename = f"duty_chart_{chart.id}_{datetime.date.today().isoformat()}.docx"
            resp = HttpResponse(
                bio.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            resp["Content-Disposition"] = f'attachment; filename="{filename}"'
            return resp

        return Response({"detail": f"Unsupported format: {out_format}"}, status=status.HTTP_400_BAD_REQUEST)


class DutyChartImportTemplateView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Download Excel template for duty chart import.",
        manual_parameters=[
            openapi.Parameter("office_id", openapi.IN_QUERY, type=openapi.TYPE_INTEGER, required=True),
            openapi.Parameter("start_date", openapi.IN_QUERY, type=openapi.TYPE_STRING, required=True),
            openapi.Parameter("end_date", openapi.IN_QUERY, type=openapi.TYPE_STRING, required=True),
            openapi.Parameter("schedule_ids", openapi.IN_QUERY, type=openapi.TYPE_ARRAY, items=openapi.Items(type=openapi.TYPE_INTEGER), required=True),
        ],
    )
    def get(self, request):
        office_id = request.query_params.get("office_id")
        start_date_str = request.query_params.get("start_date")
        end_date_str = request.query_params.get("end_date")
        # Handle both 'schedule_ids' and Axios-style 'schedule_ids[]'
        schedule_ids = request.query_params.getlist("schedule_ids") or request.query_params.getlist("schedule_ids[]")

        missing = []
        if not office_id: missing.append("office_id")
        if not start_date_str: missing.append("start_date")
        if not end_date_str: missing.append("end_date")
        if not schedule_ids: missing.append("schedule_ids")

        if missing:
            return Response({"detail": f"Missing parameters: {', '.join(missing)}"}, status=status.HTTP_400_BAD_REQUEST)

        office = get_object_or_404(Office, pk=int(office_id))
        start_date = parse_date(start_date_str)
        end_date = parse_date(end_date_str)
        schedules = Schedule.objects.filter(id__in=[int(sid) for sid in schedule_ids])

        if not (start_date and end_date):
            return Response({"detail": "Invalid dates"}, status=status.HTTP_400_BAD_REQUEST)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Duty Import Template"

        headers = ["Date (BS)", "Search Employee", "Employee ID", "Employee Name", "Phone", "Office", "Schedule", "Start Time", "End Time", "Position"]
        ws.append(headers)

        days = (end_date - start_date).days + 1
        row_idx = 2
        for sch in schedules:
            for i in range(days):
                duty_date = start_date + timedelta(days=i)
                # Duplicate rows (2 per date/schedule)
                for _ in range(2):
                    # Formulas for lookup from Reference sheet
                    # Ref Sheet Col Order: Full Name (A/1), Employee ID (B/2), Phone (C/3), Dir (D/4), Dept (E/5), Pos (F/6), Office (G/7), Search String (H/8)
                    match_val = f"MATCH(B{row_idx}, 'Reference - Office Users'!$H$2:$H$1000, 0)"
                    
                    def get_f(col_idx):
                        ref_col = f"'Reference - Office Users'!${chr(64+col_idx)}2:${chr(64+col_idx)}1000"
                        return f'=IF(B{row_idx}<>"", IFERROR(INDEX({ref_col}, {match_val}), ""), "")'

                    f_id = get_f(2)   # Employee ID
                    f_name = get_f(1) # Full Name
                    f_phone = get_f(3)
                    f_office = get_f(7) # Office
                    f_pos = get_f(6)

                    if nepali_datetime:
                        bs_date_str = nepali_datetime.date.from_datetime_date(duty_date).strftime("%Y-%m-%d")
                    else:
                        bs_date_str = duty_date.isoformat()

                    ws.append([
                        bs_date_str,
                        "",       # Search Employee (Dropdown Box)
                        f_id,     # Employee ID (Auto-populated)
                        f_name,   # Employee Name (Auto-populated)
                        f_phone,
                        f_office, # Office (Auto-populated)
                        sch.name,
                        sch.start_time.strftime("%H:%M"),
                        sch.end_time.strftime("%H:%M"),
                        f_pos
                    ])
                    row_idx += 1

        # Add Data Validation (Combined ID - Name Dropdown for Search Box)
        from openpyxl.worksheet.datavalidation import DataValidation
        dv = DataValidation(type="list", formula1="'Reference - Office Users'!$H$2:$H$1000", allow_blank=True, showErrorMessage=False)
        dv.add(f"B2:B{row_idx}")
        ws.add_data_validation(dv)
        
        # Auto-adjust column width for Search Combobox
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 25

        # Add a sheet for reference users
        # If user has 'assign_any_office_employee' permission, show all users; else only office users.
        can_assign_all = IsSuperAdmin().has_permission(request, self) or \
                         user_has_permission_slug(request.user, 'duties.assign_any_office_employee')

        ws_users = wb.create_sheet("Reference - Office Users")
        # Column Order: Full Name, Employee ID, Phone, Directorate, Department, Position, Office, ID - Name
        ws_users.append(["Full Name", "Employee ID", "Phone", "Directorate", "Department", "Position", "Office", "ID - Name"])
        from users.models import User
        if can_assign_all:
            users = User.objects.filter(is_activated=True).select_related('directorate', 'department', 'position', 'office')
        else:
            users = User.objects.filter(office=office, is_activated=True).select_related('directorate', 'department', 'position', 'office')
        
        for u in users:
            ws_users.append([
                u.full_name,
                u.employee_id,
                u.phone_number,
                getattr(u.directorate, 'name', ''),
                getattr(u.department, 'name', ''),
                getattr(u.position, 'name', ''),
                getattr(u.office, 'name', ''),
                f"{u.employee_id} - {u.full_name} ({getattr(u.office, 'name', 'N/A')})"
            ])

        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)

        filename = f"duty_template_{office.name}_{start_date_str}.xlsx"
        resp = HttpResponse(bio.getvalue(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp


class DutyChartImportView(APIView):
    permission_classes = [AdminOrReadOnly]
    parser_classes = [MultiPartParser, FormParser]

    @swagger_auto_schema(
        operation_description="Import duty chart from filled Excel template.",
        manual_parameters=[
            openapi.Parameter("office", openapi.IN_FORM, type=openapi.TYPE_INTEGER, required=True),
            openapi.Parameter("name", openapi.IN_FORM, type=openapi.TYPE_STRING, required=True),
            openapi.Parameter("effective_date", openapi.IN_FORM, type=openapi.TYPE_STRING, required=True),
            openapi.Parameter("end_date", openapi.IN_FORM, type=openapi.TYPE_STRING, required=False),
            openapi.Parameter("file", openapi.IN_FORM, type=openapi.TYPE_FILE, required=True),
        ],
    )
    def post(self, request):
        file_obj = request.FILES.get("file")
        office_id = request.data.get("office")
        name = request.data.get("name")
        effective_date_str = request.data.get("effective_date")
        end_date_str = request.data.get("end_date")
        schedule_ids = request.data.getlist("schedule_ids")
        chart_id = request.data.get("chart_id")
        dry_run = request.data.get("dry_run", "false").lower() == "true"

        if not (file_obj and office_id and effective_date_str):
            return Response({"detail": "file, office, and effective_date are required"}, status=status.HTTP_400_BAD_REQUEST)
        
        if not schedule_ids:
            return Response({"detail": "At least one schedule (shift) must be selected for the duty chart."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            df = pd.read_excel(file_obj)
            # Normalize headers
            df.columns = [str(c).strip() for c in df.columns]
        except Exception as e:
            return Response({"detail": f"Invalid Excel file: {e}"}, status=status.HTTP_400_BAD_REQUEST)

        # 0. Check for required columns (allow both "Date" and "Date (BS)")
        valid_date_headers = ["Date", "Date (BS)"]
        date_col = next((c for c in valid_date_headers if c in df.columns), None)
        
        required_cols = ["Employee ID", "Employee Name", "Schedule", "Office"]
        if not date_col or any(c not in df.columns for c in required_cols):
            missing_cols = [c for c in required_cols if c not in df.columns]
            if not date_col: missing_cols.insert(0, "Date (BS)")
            return Response({
                "detail": f"Missing required columns in Excel: {', '.join(missing_cols)}. Please use the provided template."
            }, status=status.HTTP_400_BAD_REQUEST)

        if df.empty:
            return Response({"detail": "The uploaded Excel file is empty."}, status=status.HTTP_400_BAD_REQUEST)

        office = get_object_or_404(Office, pk=int(office_id))
        
        # Permission Check: Check if user can import for this office
        user = request.user
        if not IsSuperAdmin().has_permission(request, self):
            allowed = get_allowed_office_ids(user)
            if int(office_id) not in allowed:
                if not user_has_permission_slug(user, 'duties.create_any_office_chart'):
                    return Response({
                        "detail": f"You do not have permission to import duty charts for {office.name}."
                    }, status=status.HTTP_403_FORBIDDEN)

        from django.db import transaction
        from users.models import User

        # Parse dates early for chart validation
        eff_date = parse_date(effective_date_str)
        en_date = parse_date(end_date_str) if end_date_str else None

        # 1. Duty Chart Duplicate/Retrieve Logic
        chart = None
        if chart_id:
            chart = get_object_or_404(DutyChart, pk=int(chart_id))
            # Optional: Verify office matches
            if chart.office_id != int(office_id):
                return Response({"detail": "Office mismatch for provided Duty Chart ID."}, status=status.HTTP_400_BAD_REQUEST)
        else:
            # Check for overlaps only when creating a NEW chart
            existing_charts = DutyChart.objects.filter(
                office=office,
                effective_date=eff_date,
                end_date=en_date
            )
            if schedule_ids:
                target_schedules = [int(sid) for sid in schedule_ids]
                for ec in existing_charts:
                    overlap = ec.schedules.filter(id__in=target_schedules).first()
                    if overlap:
                        return Response({
                            "detail": f"A Duty Chart already exists for '{office.name}' from {eff_date} to {en_date or 'Open'} that already includes the shift '{overlap.name}'."
                        }, status=status.HTTP_400_BAD_REQUEST)

        created_count = 0
        errors = []
        seen_in_file = set() # To detect duplicate assignments within the Excel itself

        try:
            with transaction.atomic():
                # Create Duty Chart if not provided
                if not chart:
                    chart = DutyChart.objects.create(
                        office=office,
                        name=name,
                        effective_date=eff_date,
                        end_date=en_date
                    )
                    if schedule_ids:
                        chart.schedules.set([int(sid) for sid in schedule_ids])
                else:
                    # Update metadata for existing chart
                    if name: chart.name = name
                    if eff_date: chart.effective_date = eff_date
                    chart.office = office
                    # Always update end_date (allows clearing it)
                    chart.end_date = en_date
                    chart.save()

                    # If appending, ensure schedules provided are ADDED to the chart
                    if schedule_ids:
                        new_sids = [int(sid) for sid in schedule_ids]
                        chart.schedules.add(*new_sids)

                # 2. Parse Rows and Create Duties
                today = datetime.date.today()

                preview_data = []
                for idx, row in df.iterrows():
                    row_num = idx + 2
                    row_date_val = row.get(date_col)
                    emp_id = row.get("Employee ID")
                    emp_name = row.get("Employee Name")
                    sch_name = row.get("Schedule")
                    off_name = row.get("Office")
                    
                    # If both ID and Name are missing, skip
                    if (pd.isna(emp_id) or not str(emp_id).strip()) and (pd.isna(emp_name) or not str(emp_name).strip()):
                        continue

                    try:
                        # --- A. Office Validation ---
                        # Relaxed: Only warn if office mismatch but allow if permission exists later
                        if not pd.isna(off_name) and str(off_name).strip().lower() != office.name.lower():
                            if not user_has_permission_slug(request.user, 'duties.assign_any_office_employee'):
                                errors.append(f"Row {row_num}: Office mismatch. Expected '{office.name}', found '{off_name}'. You do not have permission to assign employees from other offices.")
                                continue

                        # --- B. Date Validation ---
                        if isinstance(row_date_val, (datetime.datetime, datetime.date)):
                            temp_date = row_date_val.date() if isinstance(row_date_val, datetime.datetime) else row_date_val
                            # If year > 2070, treat as BS date incorrectly parsed by Excel/Pandas as AD
                            if temp_date.year > 2070 and nepali_datetime:
                                date_str = temp_date.strftime("%Y-%m-%d")
                                try:
                                    duty_date = nepali_datetime.datetime.strptime(date_str, "%Y-%m-%d").to_datetime_date()
                                except:
                                    duty_date = temp_date
                            else:
                                duty_date = temp_date
                        else:
                            date_str = str(row_date_val).strip()
                            duty_date = None
                            if nepali_datetime:
                                try:
                                    # Try BS parsing first
                                    if "-" in date_str:
                                        duty_date = nepali_datetime.datetime.strptime(date_str, "%Y-%m-%d").to_datetime_date()
                                    elif "/" in date_str:
                                        duty_date = nepali_datetime.datetime.strptime(date_str, "%Y/%m/%d").to_datetime_date()
                                except:
                                    pass
                            
                            # Fallback to AD parsing
                            if not duty_date:
                                duty_date = parse_date(date_str)

                        if not duty_date:
                            errors.append(f"Row {row_num}: Invalid date format '{row_date_val}'.")
                            continue
                        
                        if duty_date < today:
                            # Allow Super Admins to bypass the "past date" restriction
                            if not IsSuperAdmin().has_permission(request, self):
                                errors.append(f"Row {row_num}: Date {duty_date} is in the past. Only today or future dates allowed.")
                                continue
                        
                        if duty_date < eff_date or (en_date and duty_date > en_date):
                            error_range = f"{eff_date} to {en_date or 'Open'}"
                            errors.append(f"Row {row_num}: Date {duty_date} is outside chart range ({error_range}).")
                            continue

                        # --- C. Employee Validation ---
                        user = None
                        emp_id_str = str(emp_id).strip() if not pd.isna(emp_id) else ""
                        if " - " in emp_id_str:
                            emp_id_str = emp_id_str.split(" - ")[0].strip()

                        if emp_id_str:
                            user = User.objects.filter(employee_id__iexact=emp_id_str).first()
                        
                        if not user and not pd.isna(emp_name) and str(emp_name).strip():
                            user = User.objects.filter(full_name__iexact=str(emp_name).strip()).first() or \
                                   User.objects.filter(username__iexact=str(emp_name).strip()).first()

                        if not user:
                            errors.append(f"Row {row_num}: Employee '{emp_id or emp_name}' not found.")
                            continue

                        # --- Activation check ---
                        if not user.is_activated:
                            # Skip deactivated employees as requested
                            continue

                        # --- Office check for User (unless assign_any_office_employee permission) ---
                        user_office_id = getattr(user, 'office_id', None)
                        if user_office_id is None or int(user_office_id) != int(office_id):
                             if not user_has_permission_slug(request.user, 'duties.assign_any_office_employee'):
                                 errors.append(f"Row {row_num}: You cannot assign employee {user.full_name} as they belong to another office (or no office) and you lack the 'Assign Any Office Employee' permission.")
                                 continue

                        # --- D. Schedule & Time Validation ---
                        sch_name_str = str(sch_name).strip()
                        schedule = Schedule.objects.filter(name=sch_name_str, office=office).first() or \
                                   Schedule.objects.filter(name=sch_name_str, office__isnull=True).first()
                        
                        if not schedule:
                            errors.append(f"Row {row_num}: Schedule '{sch_name}' doesn't exist.")
                            continue

                        # Normalize and Compare Times
                        def normalize_time_str(t):
                            if isinstance(t, datetime.time): return t.strftime("%H:%M")
                            if isinstance(t, str) and len(t) >= 5: return t[:5]
                            return str(t)

                        sch_start_str = schedule.start_time.strftime("%H:%M")
                        sch_end_str = schedule.end_time.strftime("%H:%M")
                        excel_start = normalize_time_str(row.get("Start Time"))
                        excel_end = normalize_time_str(row.get("End Time"))

                        if excel_start != sch_start_str or excel_end != sch_end_str:
                            errors.append(f"Row {row_num}: Time mismatch for '{sch_name}'. Expected {sch_start_str}-{sch_end_str}, found {excel_start}-{excel_end}.")
                            continue

                        # --- E. Collision Detection (Global & Internal) ---
                        assignment_key = (user.id, duty_date, schedule.id)
                        
                        # 1. Internal check (Duplicate in file)
                        if assignment_key in seen_in_file:
                            errors.append(f"Row {row_num}: Duplicate entry for {user.full_name} on {duty_date} ({sch_name}) found within the Excel file.")
                            continue
                        seen_in_file.add(assignment_key)

                        # 2. Global check (Existing in database across all charts)
                        if Duty.objects.filter(user=user, date=duty_date, schedule=schedule).exists():
                            errors.append(f"Row {row_num}: {user.full_name} is already assigned to {sch_name} on {duty_date} in another duty chart.")
                            continue

                        # --- F. Preview Data Collect ---
                        nepali_date_str = ""
                        if nepali_datetime:
                            nepali_date_str = nepali_datetime.date.from_datetime_date(duty_date).strftime("%Y-%m-%d")

                        preview_data.append({
                            "row": row_num,
                            "date": str(duty_date),
                            "nepali_date": nepali_date_str,
                            "employee_id": user.employee_id,
                            "employee_name": user.full_name or user.username,
                            "schedule": sch_name_str,
                            "time": f"{sch_start_str} - {sch_end_str}",
                            "office": office.name
                        })

                        if not dry_run:
                            # --- G. Create Duty ---
                            Duty.objects.create(
                                user=user,
                                office=office,
                                schedule=schedule,
                                date=duty_date,
                                duty_chart=chart
                            )
                        created_count += 1
                    except Exception as e:
                        errors.append(f"Row {row_num}: Error: {str(e)}")

                if created_count == 0 and not errors:
                    errors.append("No valid duty assignments found in the Excel file. Please ensure you have selected or entered employees in the 'Employee ID' or 'Employee Name' columns.")

                if errors:
                    transaction.set_rollback(True)
                    return Response({
                        "detail": "Import failed due to validation errors. No data was saved.",
                        "errors": errors[:20]
                    }, status=status.HTTP_400_BAD_REQUEST)

                if dry_run:
                    transaction.set_rollback(True)
                    return Response({
                        "detail": "Dry run complete. No errors found.",
                        "is_preview": True,
                        "created_duties": created_count,
                        "preview_data": preview_data
                    }, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"detail": f"An unexpected error occurred: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response({
            "detail": "Import complete",
            "chart_id": chart.id,
            "effective_date": chart.effective_date,
            "created_duties": created_count
        }, status=status.HTTP_201_CREATED)

